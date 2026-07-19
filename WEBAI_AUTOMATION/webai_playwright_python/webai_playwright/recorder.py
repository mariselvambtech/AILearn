from __future__ import annotations

import asyncio
import time
from dataclasses import dataclass, asdict, field
from typing import Any, Dict, List, Optional
import json

from playwright.async_api import Page


@dataclass
class Step:
    action: str  # open | navigate | click | type | press_key | verify_text | verify_visible | extract | extract_table
    url: Optional[str] = None
    name: Optional[str] = None     # click name OR confirmed label anchor OR key press context OR extraction variable name
    value: Optional[str] = None    # typed value OR verify text OR extraction sample value
    key: Optional[str] = None      # pressed key (for press_key action) OR extract_type ('text'/'attribute')
    ts: float = 0.0
    locators: Optional[List[Dict[str, str]]] = None  # NEW: Multiple locator strategies
    extract_type: Optional[str] = None  # NEW: For extract action - 'text' or 'attribute'
    attribute_name: Optional[str] = None  # NEW: For extract action with type='attribute'
    save_options: Optional[Dict[str, Any]] = None  # NEW: Phase 8.2 - Save configuration for extract action
    table_config: Optional[Dict[str, Any]] = None  # NEW: Phase 8.3 - Table extraction configuration


RECORDER_INIT_SCRIPT = r"""
(() => {
  let __RECORDER_STOPPED__ = false;
  let verifyVisibleArmed = false;

  const lastSentValue = new WeakMap();
  const suggestedLabel = new WeakMap();
  const labelDecision = new WeakMap();
  const userInteracted = new WeakMap();  // NEW: Track user interactions

  const MAX_CLICK_TEXT = 60;
  const MAX_LABEL_TEXT = 60;

  function isRecorderUi(el) {
    if (!el || !el.closest) return false;
    return Boolean(
      el.closest("#__webai_recorder_modal__") ||
      el.closest("#__webai_recorder_hint__") ||
      el.closest("#__webai_recorder_stopped__") ||
      el.closest("#__webai_recorder_stopbtn__") ||
      el.closest("#__extraction_menu__") ||          // Extraction menu
      el.closest("#__extraction_input__") ||         // Extraction input dialog
      el.closest("[data-webai-dialog='true']")       // NEW: Save dialogs
    );
  }

  function isVisible(el) {
    if (!el) return false;
    const r = el.getBoundingClientRect();
    if (r.width < 2 || r.height < 2) return false;
    const style = window.getComputedStyle(el);
    if (style.visibility === "hidden" || style.display === "none" || style.opacity === "0") return false;
    return true;
  }

  function cleanText(s) {
    return (s || "").replace(/\s+/g, " ").trim();
  }

  function truncate(s, n) {
    s = cleanText(s);
    if (!s) return "";
    return s.length <= n ? s : s.slice(0, n) + "…";
  }

  // Prefer "stable" click targets only
  function getClickableTarget(el) {
    if (!el) return null;
    // true interactives
    const target = el.closest?.("a,button,[role='button'],input[type='submit'],input[type='button']");
    return target || null;
  }

  function bestStableName(el, maxLen) {
    if (!el) return null;

    const aria = el.getAttribute && el.getAttribute("aria-label");
    if (aria && cleanText(aria)) return truncate(aria, maxLen);

    const title = el.getAttribute && el.getAttribute("title");
    if (title && cleanText(title)) return truncate(title, maxLen);

    if (el.id) {
      const lbl = document.querySelector(`label[for="${CSS.escape(el.id)}"]`);
      if (lbl && isVisible(lbl)) {
        const t = cleanText(lbl.innerText);
        if (t) return truncate(t, maxLen);
      }
    }

    const ph = el.getAttribute && el.getAttribute("placeholder");
    if (ph && cleanText(ph)) return truncate(ph, maxLen);

    // for buttons/links: use their visible text (trimmed)
    const txt = cleanText(el.innerText || "");
    if (txt) return truncate(txt, maxLen);

    const nm = el.getAttribute && el.getAttribute("name");
    if (nm && cleanText(nm)) return truncate(nm, maxLen);
    
    // NEW: Check for alt attribute (images)
    const alt = el.getAttribute && el.getAttribute("alt");
    if (alt && cleanText(alt)) return truncate(alt, maxLen);

    // NEW: Fallback for icon buttons - use contextual detection
    const tag = el.tagName ? el.tagName.toLowerCase() : "";
    const type = el.getAttribute && el.getAttribute("type");
    const role = el.getAttribute && el.getAttribute("role");
    
    // Special handling for search buttons
    if (tag === "button" || role === "button" || type === "submit") {
      // Check if it's near a search input
      const parent = el.closest("form") || el.parentElement;
      if (parent) {
        const searchInput = parent.querySelector("input[type='search']") || 
                          parent.querySelector("input[placeholder*='search' i]") ||
                          parent.querySelector("input[placeholder*='Search' i]") ||
                          parent.querySelector("input[aria-label*='search' i]") ||
                          parent.querySelector("input[aria-label*='Search' i]");
        if (searchInput) {
          return "Search";
        }
      }
      
      // Generic button fallback
      return "Button";
    }
    
    // Link fallback
    if (tag === "a") {
      return "Link";
    }

    // Last resort: return null to filter out truly meaningless clicks
    return null;
  }

  // ============================================================================
  // NEW: Multi-Locator Strategy System
  // ============================================================================

  /**
   * Get implicit ARIA role for common HTML elements
   */
  function getImplicitRole(el) {
    if (!el || !el.tagName) return null;
    const tag = el.tagName.toLowerCase();
    const type = el.getAttribute && el.getAttribute('type');
    
    const roleMap = {
      'button': 'button',
      'a': 'link',
      'img': 'img',
      'input': type === 'text' ? 'textbox' : type === 'submit' ? 'button' : 'textbox',
      'textarea': 'textbox',
      'select': 'combobox',
      'nav': 'navigation',
      'header': 'banner',
      'footer': 'contentinfo',
      'main': 'main',
      'h1': 'heading',
      'h2': 'heading',
      'h3': 'heading',
      'h4': 'heading',
      'h5': 'heading',
      'h6': 'heading',
    };
    
    return roleMap[tag] || null;
  }

  /**
   * Generate stable CSS selector for element
   */
  function generateStableCSS(el) {
    if (!el || !el.tagName) return null;
    
    const parts = [];
    const tag = el.tagName.toLowerCase();
    parts.push(tag);
    
    // Add ID if present (makes it unique)
    if (el.id) {
      return `#${CSS.escape(el.id)}`;
    }
    
    // Add stable classes (exclude dynamic ones like 'active', 'selected', etc.)
    const dynamicClassPatterns = /active|selected|hover|focus|disabled|loading|open|closed|visible|hidden/i;
    const classes = Array.from(el.classList || [])
      .filter(c => !dynamicClassPatterns.test(c))
      .slice(0, 2);
    
    if (classes.length > 0) {
      parts[0] += '.' + classes.map(c => CSS.escape(c)).join('.');
    }
    
    // Add attribute selectors for more specificity
    const stableAttrs = ['type', 'name', 'data-testid'];
    for (const attr of stableAttrs) {
      const val = el.getAttribute(attr);
      if (val) {
        parts[0] += `[${attr}="${CSS.escape(val)}"]`;
        break; // One attribute is usually enough
      }
    }
    
    const selector = parts.join('');
    
    // Check if selector is unique
    try {
      const matches = document.querySelectorAll(selector);
      if (matches.length === 1) {
        return selector;
      }
      
      // Add nth-child if not unique
      let index = 1;
      let sibling = el.previousElementSibling;
      while (sibling) {
        if (sibling.tagName === el.tagName) index++;
        sibling = sibling.previousElementSibling;
      }
      
      return `${selector}:nth-child(${index})`;
    } catch (e) {
      return selector;
    }
  }

  /**
   * Generate XPath for element (last resort locator)
   */
  function generateXPath(el) {
    if (!el || el.nodeType !== 1) return null;
    
    // If element has ID, use it (most stable XPath)
    if (el.id) {
      return `//*[@id="${el.id}"]`;
    }
    
    const parts = [];
    let current = el;
    let depth = 0;
    
    while (current && current.nodeType === Node.ELEMENT_NODE && current.tagName !== 'BODY') {
      const tag = current.nodeName.toLowerCase();
      
      // Count preceding siblings of same tag
      let index = 1;
      let sibling = current.previousSibling;
      while (sibling) {
        if (sibling.nodeType === Node.ELEMENT_NODE && sibling.nodeName === current.nodeName) {
          index++;
        }
        sibling = sibling.previousSibling;
      }
      
      // Check if index is needed (i.e., are there siblings of same tag?)
      const totalSiblings = Array.from(current.parentNode?.children || [])
        .filter(s => s.nodeName === current.nodeName).length;
      
      const part = totalSiblings > 1 ? `${tag}[${index}]` : tag;
      parts.unshift(part);
      
      current = current.parentElement;
      depth++;
      
      // Limit depth to keep XPath reasonable
      if (depth >= 5) break;
    }
    
    return '//' + parts.join('/');
  }

  /**
   * Collect all available locator strategies for an element
   * Returns array of {type, value} objects ordered by stability
   */
  function getLocatorCandidates(el) {
    if (!el) return [];
    
    const candidates = [];
    
    // 1. Test ID (HIGHEST PRIORITY - developer-intended selector)
    const testId = el.getAttribute && (
      el.getAttribute('data-testid') ||
      el.getAttribute('data-test-id') ||
      el.getAttribute('data-test')
    );
    if (testId) {
      candidates.push({type: 'test-id', value: testId});
    }
    
    // 2. ID (very stable if present)
    if (el.id && !/^[0-9]/.test(el.id)) {  // Exclude numeric IDs (often auto-generated)
      candidates.push({type: 'id', value: el.id});
    }
    
    // 3. Name attribute (stable for form elements)
    const name = el.getAttribute && el.getAttribute('name');
    if (name) {
      candidates.push({type: 'name', value: name});
    }
    
    // 4. Aria-label (semantic and accessible)
    const ariaLabel = el.getAttribute && el.getAttribute('aria-label');
    if (ariaLabel && cleanText(ariaLabel)) {
      candidates.push({type: 'aria-label', value: cleanText(ariaLabel)});
    }
    
    // 5. Placeholder (for input fields)
    const placeholder = el.getAttribute && el.getAttribute('placeholder');
    if (placeholder && cleanText(placeholder)) {
      candidates.push({type: 'placeholder', value: cleanText(placeholder)});
    }
    
    // 5a. Title attribute (tooltips, accessible names)
    const title = el.getAttribute && el.getAttribute('title');
    if (title && cleanText(title)) {
      candidates.push({type: 'title', value: cleanText(title)});
    }
    
    // 5b. Alt attribute (images, icons)
    const alt = el.getAttribute && el.getAttribute('alt');
    if (alt && cleanText(alt)) {
      candidates.push({type: 'alt', value: cleanText(alt)});
    }
    
    // 5c. Href attribute (links - very stable for navigation)
    if (el.tagName && el.tagName.toLowerCase() === 'a') {
      const href = el.getAttribute && el.getAttribute('href');
      if (href && href.trim() && !href.startsWith('javascript:')) {
        candidates.push({type: 'href', value: href.trim()});
      }
    }
    
    // 6. Label (find associated label element)
    if (el.id) {
      const label = document.querySelector(`label[for="${CSS.escape(el.id)}"]`);
      if (label) {
        const labelText = cleanText(label.innerText);
        if (labelText) {
          candidates.push({type: 'label', value: labelText});
        }
      }
    }
    
    // 7. CSS Selector (powerful and specific)
    const css = generateStableCSS(el);
    if (css) {
      candidates.push({type: 'css', value: css});
    }
    
    // 8. Text content (brittle - changes with content)
    const text = cleanText(el.innerText || el.textContent || '');
    if (text && text.length < 50) {  // Avoid huge text blocks
      candidates.push({type: 'text', value: text});
    }
    
    // 9. Role + Name (for semantic elements)
    const explicitRole = el.getAttribute && el.getAttribute('role');
    const role = explicitRole || getImplicitRole(el);
    if (role) {
      candidates.push({
        type: 'role',
        value: role,
        name: text || ariaLabel || null
      });
    }
    
    // 10. XPath (LAST RESORT - most brittle)
    const xpath = generateXPath(el);
    if (xpath) {
      candidates.push({type: 'xpath', value: xpath});
    }
    
    return candidates;
  }

  // ============================================================================
  // End of Multi-Locator System
  // ============================================================================

  function send(kind, payload) {
    if (__RECORDER_STOPPED__) return;
    try { if (window.__recordEvent) window.__recordEvent({ kind, ...payload }); } catch {}
  }

  function showHint(msg, id="__webai_recorder_hint__") {
    const old = document.getElementById(id);
    if (old) old.remove();
    const d = document.createElement("div");
    d.id = id;
    d.textContent = msg;
    d.style.position = "fixed";
    d.style.bottom = "15px";
    d.style.left = "15px";
    d.style.background = "rgba(0,0,0,0.85)";
    d.style.color = "white";
    d.style.padding = "8px 12px";
    d.style.borderRadius = "8px";
    d.style.zIndex = "2147483647";
    d.style.fontFamily = "Arial, sans-serif";
    document.body.appendChild(d);
    setTimeout(() => d.remove(), 2500);
  }

  function showPrompt(title, placeholder) {
    return new Promise((resolve) => {
      const old = document.getElementById("__webai_recorder_modal__");
      if (old) old.remove();

      const wrap = document.createElement("div");
      wrap.id = "__webai_recorder_modal__";
      wrap.style.position = "fixed";
      wrap.style.inset = "0";
      wrap.style.background = "rgba(0,0,0,0.45)";
      wrap.style.zIndex = "2147483647";
      wrap.style.display = "flex";
      wrap.style.alignItems = "center";
      wrap.style.justifyContent = "center";

      const box = document.createElement("div");
      box.style.background = "white";
      box.style.padding = "16px";
      box.style.borderRadius = "10px";
      box.style.width = "540px";
      box.style.maxWidth = "92vw";
      box.style.fontFamily = "Arial, sans-serif";
      box.style.boxShadow = "0 10px 30px rgba(0,0,0,0.35)";

      const h = document.createElement("div");
      h.textContent = title;
      h.style.fontWeight = "700";
      h.style.marginBottom = "8px";

      const input = document.createElement("input");
      input.placeholder = placeholder || "";
      input.style.width = "100%";
      input.style.padding = "10px";
      input.style.border = "1px solid #ccc";
      input.style.borderRadius = "8px";
      input.style.outline = "none";

      const row = document.createElement("div");
      row.style.display = "flex";
      row.style.justifyContent = "flex-end";
      row.style.gap = "10px";
      row.style.marginTop = "12px";

      const cancel = document.createElement("button");
      cancel.textContent = "Cancel";
      cancel.style.padding = "8px 12px";
      cancel.style.borderRadius = "8px";
      cancel.style.border = "1px solid #ccc";
      cancel.style.background = "#f7f7f7";
      cancel.style.cursor = "pointer";

      const ok = document.createElement("button");
      ok.textContent = "OK";
      ok.style.padding = "8px 12px";
      ok.style.borderRadius = "8px";
      ok.style.border = "1px solid #0a66ff";
      ok.style.background = "#0a66ff";
      ok.style.color = "white";
      ok.style.cursor = "pointer";

      function cleanup(val) { wrap.remove(); resolve(val); }

      cancel.onclick = () => cleanup(null);
      ok.onclick = () => cleanup(cleanText(input.value) || null);

      wrap.addEventListener("keydown", (e) => {
        if (e.key === "Escape") { e.preventDefault(); e.stopPropagation(); cleanup(null); }
        if (e.key === "Enter")  { e.preventDefault(); e.stopPropagation(); cleanup(cleanText(input.value) || null); }
      }, true);

      row.appendChild(cancel);
      row.appendChild(ok);

      box.appendChild(h);
      box.appendChild(input);
      box.appendChild(row);

      wrap.appendChild(box);
      document.body.appendChild(wrap);
      setTimeout(() => input.focus(), 0);
    });
  }

  function confirmLabel(suggested) {
    return new Promise((resolve) => {
      const old = document.getElementById("__webai_recorder_modal__");
      if (old) old.remove();

      const wrap = document.createElement("div");
      wrap.id = "__webai_recorder_modal__";
      wrap.style.position = "fixed";
      wrap.style.inset = "0";
      wrap.style.background = "rgba(0,0,0,0.45)";
      wrap.style.zIndex = "2147483647";
      wrap.style.display = "flex";
      wrap.style.alignItems = "center";
      wrap.style.justifyContent = "center";

      const box = document.createElement("div");
      box.style.background = "white";
      box.style.padding = "16px";
      box.style.borderRadius = "10px";
      box.style.width = "600px";
      box.style.maxWidth = "92vw";
      box.style.fontFamily = "Arial, sans-serif";
      box.style.boxShadow = "0 10px 30px rgba(0,0,0,0.35)";

      const h = document.createElement("div");
      h.textContent = "Anchor/Label detected for this input";
      h.style.fontWeight = "700";
      h.style.marginBottom = "8px";

      const p = document.createElement("div");
      p.textContent = `Suggested label: "${suggested}"`;
      p.style.marginBottom = "12px";

      const row = document.createElement("div");
      row.style.display = "flex";
      row.style.justifyContent = "flex-end";
      row.style.gap = "10px";

      const useBtn = document.createElement("button");
      useBtn.textContent = "Use label";
      useBtn.style.padding = "8px 12px";
      useBtn.style.borderRadius = "8px";
      useBtn.style.border = "1px solid #0a66ff";
      useBtn.style.background = "#0a66ff";
      useBtn.style.color = "white";
      useBtn.style.cursor = "pointer";

      const skipBtn = document.createElement("button");
      skipBtn.textContent = "Skip";
      skipBtn.style.padding = "8px 12px";
      skipBtn.style.borderRadius = "8px";
      skipBtn.style.border = "1px solid #ccc";
      skipBtn.style.background = "#f7f7f7";
      skipBtn.style.cursor = "pointer";

      const editBtn = document.createElement("button");
      editBtn.textContent = "Edit…";
      editBtn.style.padding = "8px 12px";
      editBtn.style.borderRadius = "8px";
      editBtn.style.border = "1px solid #ccc";
      editBtn.style.background = "#f7f7f7";
      editBtn.style.cursor = "pointer";

      function cleanup(val) { wrap.remove(); resolve(val); }

      useBtn.onclick = () => cleanup(suggested);
      skipBtn.onclick = () => cleanup(null);
      editBtn.onclick = async () => {
        wrap.remove();
        const custom = await showPrompt("Enter label to use for this input", suggested);
        resolve(custom || null);
      };

      row.appendChild(skipBtn);
      row.appendChild(editBtn);
      row.appendChild(useBtn);

      box.appendChild(h);
      box.appendChild(p);
      box.appendChild(row);

      wrap.appendChild(box);
      document.body.appendChild(wrap);
    });
  }

  function findNearbyLabel(inputEl) {
    if (!inputEl) return null;

    // best stable names for inputs
    const direct = bestStableName(inputEl, MAX_LABEL_TEXT);
    if (direct) return direct;

    const r = inputEl.getBoundingClientRect();
    const cx = r.left + r.width / 2;
    const cy = r.top + r.height / 2;

    const candidates = Array.from(document.querySelectorAll("label, span, div, p, strong, b, small"));

    let best = null;
    let bestScore = 1e9;

    for (const el of candidates) {
      if (!el || el === inputEl) continue;
      if (!isVisible(el)) continue;
      if (isRecorderUi(el)) continue;

      const t = cleanText(el.innerText || "");
      if (!t) continue;
      if (t.length < 2 || t.length > 40) continue;

      const rr = el.getBoundingClientRect();
      const ex = rr.left + rr.width / 2;
      const ey = rr.top + rr.height / 2;

      const dx = ex - cx;
      const dy = ey - cy;
      const dist = Math.sqrt(dx*dx + dy*dy);

      if (dist > 220) continue;

      let penalty = 0;
      if (ey > cy + 10) penalty += 50;
      if (ex > cx + 10) penalty += 30;

      const score = dist + penalty;
      if (score < bestScore) {
        bestScore = score;
        best = t;
      }
    }
    return best ? truncate(best, MAX_LABEL_TEXT) : null;
  }

  async function finalizeTypeFor(el) {
    if (__RECORDER_STOPPED__) return;
    if (!el || isRecorderUi(el)) return;
    if (!("value" in el)) return;

    const value = cleanText(String(el.value ?? ""));
    if (!value) return; // skip empty values

    const last = lastSentValue.get(el);
    if (value === last) return;
    
    // NEW: Only record if user actually interacted with this element
    const hasInteraction = userInteracted.has(el);
    console.log(`[finalizeTypeFor] value="${value}" hasInteraction=${hasInteraction} url=${location.href}`);
    
    if (!hasInteraction) {
      console.log('  ⏭️  Skipping auto-populated value:', value);
      return;
    }

    lastSentValue.set(el, value);

    let label = labelDecision.get(el);
    if (label === undefined) {
      const suggested = suggestedLabel.get(el);
      // AUTO-ACCEPT: Use suggested label directly instead of asking for confirmation
      // This prevents typing from being lost when user quickly presses Enter
      label = suggested || "(unknown)";
      labelDecision.set(el, label);
    }

    // NEW: Collect multiple locator strategies for the input element
    const locators = getLocatorCandidates(el);

    console.log(`  ✅ Recording type: "${value}" into "${label}"`);
    send("type_final", { 
      url: location.href, 
      value, 
      label,       // Keep for backward compatibility
      locators     // NEW: Multiple locator candidates
    });
  }

  document.addEventListener("focusin", (e) => {
    const el = e.target;
    if (!el || isRecorderUi(el)) return;
    const tag = (el.tagName || "").toLowerCase();
    if (!(tag === "input" || tag === "textarea")) return;
    if (!suggestedLabel.has(el)) {
      suggestedLabel.set(el, findNearbyLabel(el));
    }
    // NOTE: We do NOT mark as interacted here! 
    // Clicking to select text ≠ typing
  }, true);

  // NEW: Also track when user types
  document.addEventListener("input", (e) => {
    const el = e.target;
    if (!el || isRecorderUi(el)) return;
    const tag = (el.tagName || "").toLowerCase();
    if (!(tag === "input" || tag === "textarea")) return;
    // Mark that user typed in this element
    console.log(`[input] Marking element as interacted: ${el.tagName} "${el.value || ''}"`);
    userInteracted.set(el, true);
  }, true);

  document.addEventListener("focusout", (e) => {
    const el = e.target;
    if (!el || isRecorderUi(el)) return;
    const tag = (el.tagName || "").toLowerCase();
    if (!(tag === "input" || tag === "textarea")) return;
    finalizeTypeFor(el);
  }, true);

  document.addEventListener("keydown", (e) => {
    if (__RECORDER_STOPPED__) return;
    
    const el = e.target;
    if (!el || isRecorderUi(el)) return;
    
    // Finalize typing on Tab/Enter
    if (e.key === "Tab" || e.key === "Enter") {
      const tag = (el.tagName || "").toLowerCase();
      if (tag === "input" || tag === "textarea") finalizeTypeFor(el);
    }
    
    // Record important key presses
    const RECORDABLE_KEYS = ["Enter", "Escape", "Tab", "ArrowUp", "ArrowDown", "ArrowLeft", "ArrowRight"];
    if (RECORDABLE_KEYS.includes(e.key)) {
      const context = getLabel(el) || null;
      actions.push({
        action: "press_key",
        url: window.location.href,
        key: e.key,
        name: context,
        value: null,
        ts: Date.now() / 1000
      });
      notify();
      
      // Visual feedback
      if (feedbackDiv) {
        feedbackDiv.textContent = `⌨️ Key: ${e.key}${context ? ` (in "${context}")` : ""}`;
        feedbackDiv.style.display = "block";
        if (feedbackTimeout) clearTimeout(feedbackTimeout);
        feedbackTimeout = setTimeout(() => {
          if (feedbackDiv) feedbackDiv.style.display = "none";
        }, 1500);
      }
    }
  }, true);

  // Stable click recording: only record if we find a real clickable target
  document.addEventListener("click", (e) => {
    if (__RECORDER_STOPPED__) return;

    const raw = e.target;
    if (!raw || isRecorderUi(raw)) return;

    if (verifyVisibleArmed) {
      const t = cleanText(raw.innerText || "");
      if (t) {
        try { window.__recordVerifyVisible({ url: location.href, text: truncate(t, MAX_CLICK_TEXT) }); } catch {}
      }
      verifyVisibleArmed = false;
      return;
    }

    const clickable = getClickableTarget(raw);
    if (!clickable) return; // <-- stability improvement: ignore random container clicks

    const name = bestStableName(clickable, MAX_CLICK_TEXT);
    if (!name) return; // <-- This was the issue for icon buttons

    // skip huge/garbage even after truncate check
    if (cleanText(name).length > MAX_CLICK_TEXT) return;

    // NEW: Collect multiple locator strategies
    const locators = getLocatorCandidates(clickable);

    send("click", { 
      url: location.href, 
      name,        // Keep for backward compatibility
      locators     // NEW: Multiple locator candidates
    });
  }, true);

  async function stopNow() {
    if (__RECORDER_STOPPED__) return;
    const active = document.activeElement;
    if (active && active.tagName) {
      const t = active.tagName.toLowerCase();
      if (t === "input" || t === "textarea") await finalizeTypeFor(active);
    }
    __RECORDER_STOPPED__ = true;
    showHint("Recorder: Stopping…", "__webai_recorder_stopped__");
    try { if (window.__stopRecording) await window.__stopRecording(); } catch {}
  }

  function ensureStopButton() {
    if (document.getElementById("__webai_recorder_stopbtn__")) return;
    const btn = document.createElement("button");
    btn.id = "__webai_recorder_stopbtn__";
    btn.textContent = "Stop Recording";
    btn.style.position = "fixed";
    btn.style.right = "16px";
    btn.style.bottom = "16px";
    btn.style.zIndex = "2147483647";
    btn.style.padding = "10px 12px";
    btn.style.borderRadius = "10px";
    btn.style.border = "1px solid #c00";
    btn.style.background = "#fff";
    btn.style.color = "#c00";
    btn.style.fontFamily = "Arial, sans-serif";
    btn.style.cursor = "pointer";
    btn.onclick = () => stopNow();
    document.body.appendChild(btn);
  }

  if (document.readyState === "loading") {
    document.addEventListener("DOMContentLoaded", ensureStopButton, { once: true });
  } else {
    ensureStopButton();
  }

  // shortcuts: Ctrl+Shift+S OR Ctrl+Alt+S OR Esc
  document.addEventListener("keydown", async (e) => {
    if (__RECORDER_STOPPED__) return;

    const stopCombo = (e.ctrlKey && e.shiftKey && e.code === "KeyS") ||
                      (e.ctrlKey && e.altKey && e.code === "KeyS") ||
                      (e.code === "Escape");
    if (stopCombo) {
      e.preventDefault();
      e.stopPropagation();
      await stopNow();
      return;
    }

    // verify text: Ctrl+Shift+V
    if (e.ctrlKey && e.shiftKey && e.code === "KeyV") {
      e.preventDefault();
      e.stopPropagation();
      const t = await showPrompt("Verify page contains text", "Example: Get in touch");
      if (t) {
        try { window.__recordVerifyText({ url: location.href, text: t }); } catch {}
      }
      return;
    }

    // verify element visible: Ctrl+Shift+E
    if (e.ctrlKey && e.shiftKey && e.code === "KeyE") {
      e.preventDefault();
      e.stopPropagation();
      verifyVisibleArmed = true;
      showHint("Recorder: Click element to verify visible");
      return;
    }

    // NEW: Add delay: Ctrl+Shift+W
    if (e.ctrlKey && e.shiftKey && e.code === "KeyW") {
      e.preventDefault();
      e.stopPropagation();
      window.__addDelay();
      return;
    }
  }, true);

  // ============================================================================
  // TABLE EXTRACTION SYSTEM - Phase 8.3
  // ============================================================================
  
  function detectTableElement(target) {
    // Walk up DOM to find <table> or container with table-like structure
    let el = target;
    while (el && el !== document.body) {
      if (el.tagName === 'TABLE') return el;
      if (el.querySelector('table')) return el.querySelector('table');
      // Check for div-based tables (e.g., role="table")
      if (el.getAttribute && el.getAttribute('role') === 'table') return el;
      el = el.parentElement;
    }
    return null;
  }
  
  function getTableHeaders(table) {
    // Try thead first
    let headers = Array.from(table.querySelectorAll('thead th, thead td'))
      .map(h => h.innerText.trim());
    
    if (headers.length === 0) {
      // Fallback: first row
      const firstRow = table.querySelector('tr');
      if (firstRow) {
        headers = Array.from(firstRow.querySelectorAll('th, td'))
          .map(h => h.innerText.trim());
      }
    }
    
    return headers.filter(h => h.length > 0);  // Remove empty headers
  }
  
  function showColumnSelectionDialog(table, callback) {
    const headers = getTableHeaders(table);
    
    if (headers.length === 0) {
      alert('⚠️ No table headers found!');
      callback(null);
      return;
    }
    
    const overlay = document.createElement('div');
    overlay.style.cssText = `
      position: fixed; top: 0; left: 0; right: 0; bottom: 0;
      background: rgba(0,0,0,0.5); z-index: 999999;
      display: flex; align-items: center; justify-content: center;
    `;
    overlay.setAttribute('data-webai-dialog', 'true');
    
    const dialog = document.createElement('div');
    dialog.style.cssText = `
      background: white; padding: 20px; border-radius: 8px;
      box-shadow: 0 4px 12px rgba(0,0,0,0.3);
      min-width: 400px; max-width: 600px; max-height: 80vh;
      overflow-y: auto;
    `;
    
    dialog.innerHTML = `
      <div style="font-weight: bold; margin-bottom: 15px; font-size: 16px;">
        📊 Select Columns to Extract
      </div>
      <div style="margin-bottom: 10px;">
        <button id="__select_all__" style="margin-right: 8px; padding: 4px 8px; cursor: pointer;">Select All</button>
        <button id="__deselect_all__" style="padding: 4px 8px; cursor: pointer;">Deselect All</button>
      </div>
      <div id="__column_list__" style="margin-bottom: 15px; max-height: 300px; overflow-y: auto; border: 1px solid #ddd; padding: 10px; border-radius: 4px;">
        ${headers.map((h, i) => `
          <label style="display: block; margin: 8px 0; cursor: pointer;">
            <input type="checkbox" class="column-checkbox" data-index="${i}" 
                   value="${h}" style="margin-right: 8px;" checked>
            <span>${h || ('Column ' + (i + 1))}</span>
          </label>
        `).join('')}
      </div>
      <div style="margin-bottom: 15px; padding: 10px; background: #f5f5f5; border-radius: 4px;">
        <label style="cursor: pointer; display: flex; align-items: center; margin-bottom: 8px;">
          <input type="checkbox" id="__pagination_enabled__" style="margin-right: 8px;">
          <span style="font-weight: bold;">Enable Pagination</span>
        </label>
        <div id="__pagination_options__" style="display: none; margin-left: 24px;">
          <label style="display: block; margin: 4px 0;">
            <span>Max Pages (1-100):</span>
            <input type="number" id="__max_pages__" value="10" min="1" max="100" 
                   style="width: 60px; margin-left: 8px; padding: 2px;">
          </label>
          <label style="display: block; margin: 4px 0;">
            <span>Wait Per Page (1-10 sec):</span>
            <input type="number" id="__wait_per_page__" value="2" min="1" max="10" step="0.5"
                   style="width: 60px; margin-left: 8px; padding: 2px;">
          </label>
          <label style="display: block; margin: 4px 0;">
            <span>Page Timeout (5-30 sec):</span>
            <input type="number" id="__page_timeout__" value="10" min="5" max="30" step="1"
                   style="width: 60px; margin-left: 8px; padding: 2px;">
          </label>
          <label style="display: block; margin: 4px 0;">
            <span>Retry Attempts (1-5):</span>
            <input type="number" id="__retry_attempts__" value="3" min="1" max="5" step="1"
                   style="width: 60px; margin-left: 8px; padding: 2px;">
          </label>
        </div>
      </div>
      <div style="display: flex; justify-content: flex-end; gap: 8px;">
        <button id="__column_cancel__" style="padding: 8px 16px; background: #999; color: white; border: none; border-radius: 4px; cursor: pointer;">Cancel</button>
        <button id="__column_extract__" style="padding: 8px 16px; background: #4CAF50; color: white; border: none; border-radius: 4px; cursor: pointer;">Extract</button>
      </div>
    `;
    
    overlay.appendChild(dialog);
    document.body.appendChild(overlay);
    
    const cleanup = () => {
      if (overlay.parentElement) {
        overlay.parentElement.removeChild(overlay);
      }
    };
    
    // Select All button
    document.getElementById('__select_all__').onclick = () => {
      document.querySelectorAll('.column-checkbox').forEach(cb => cb.checked = true);
    };
    
    // Deselect All button
    document.getElementById('__deselect_all__').onclick = () => {
      document.querySelectorAll('.column-checkbox').forEach(cb => cb.checked = false);
    };
    
    // Pagination toggle
    const paginationCheckbox = document.getElementById('__pagination_enabled__');
    const paginationOptions = document.getElementById('__pagination_options__');
    paginationCheckbox.onchange = () => {
      paginationOptions.style.display = paginationCheckbox.checked ? 'block' : 'none';
    };
    
    // Cancel button
    document.getElementById('__column_cancel__').onclick = () => {
      cleanup();
      callback(null);
    };
    
    // Extract button
    document.getElementById('__column_extract__').onclick = () => {
      const selectedColumns = [];
      const selectedIndices = [];
      
      document.querySelectorAll('.column-checkbox:checked').forEach(cb => {
        selectedColumns.push(cb.value);
        selectedIndices.push(parseInt(cb.getAttribute('data-index')));
      });
      
      if (selectedColumns.length === 0) {
        alert('⚠️ Please select at least one column!');
        return;
      }
      
      const config = {
        columns: selectedColumns,
        column_indices: selectedIndices,
        pagination: {
          enabled: paginationCheckbox.checked,
          max_pages: parseInt(document.getElementById('__max_pages__').value) || 10,
          wait_per_page: parseFloat(document.getElementById('__wait_per_page__').value) || 2.0,
          page_timeout: parseFloat(document.getElementById('__page_timeout__').value) || 10.0,
          retry_attempts: parseInt(document.getElementById('__retry_attempts__').value) || 3
        }
      };
      
      cleanup();
      callback(config);
    };
  }

  // ============================================================================
  // DATA EXTRACTION SYSTEM - Right-Click Context Menu
  // ============================================================================
  
  let __CONTEXT_MENU__ = null;
  let __CONTEXT_TARGET__ = null;

  // Handle right-click to show extraction menu
  document.addEventListener('contextmenu', (e) => {
    if (__RECORDER_STOPPED__) return;  // Only block if stopped
    
    e.preventDefault();
    e.stopPropagation();
    
    __CONTEXT_TARGET__ = e.target;
    showExtractionMenu(e.pageX, e.pageY);
  }, true);

  function showExtractionMenu(x, y) {
    removeExtractionMenu();
    
    const menu = document.createElement('div');
    menu.id = '__extraction_menu__';
    menu.style.cssText = `
      position: fixed;
      left: ${x}px;
      top: ${y}px;
      z-index: 2147483647;
      background: white;
      border: 2px solid #4CAF50;
      border-radius: 8px;
      padding: 12px;
      box-shadow: 0 4px 12px rgba(0,0,0,0.3);
      font-family: Arial, sans-serif;
      font-size: 14px;
    `;
    
    // Create buttons with event listeners instead of inline onclick
    const title = document.createElement('div');
    title.style.cssText = 'font-weight: bold; margin-bottom: 10px; color: #333; font-size: 15px;';
    title.textContent = '🔍 Extract Data';
    
    const btnText = createMenuButton('📝 Extract Text', '#4CAF50');
    btnText.addEventListener('click', (e) => {
      e.stopPropagation();
      console.log('Extract Text clicked');
      window.__extractText();
    });
    
    const btnAttr = createMenuButton('🔗 Extract Attribute', '#2196F3');
    btnAttr.addEventListener('click', (e) => {
      e.stopPropagation();
      console.log('Extract Attribute clicked');
      window.__extractAttribute();
    });
    
    const btnTable = createMenuButton('📊 Extract Table', '#FF9800');
    btnTable.addEventListener('click', (e) => {
      e.stopPropagation();
      console.log('Extract Table clicked');
      window.__extractTable();
    });
    
    const btnDelay = createMenuButton ('⏱️ Add Delay', '#9C27B0');
    btnDelay.addEventListener('click', (e) => {
      e.stopPropagation();
      console.log('Add Delay clicked');
      window.__addDelay();
    });
    
    const btnCancel = createMenuButton('❌ Cancel', '#f44336');
    btnCancel.addEventListener('click', (e) => {
      e.stopPropagation();
      console.log('Cancel clicked');
      window.__cancelExtract();
    });
    
    menu.appendChild(title);
    menu.appendChild(btnText);
    menu.appendChild(btnAttr);
    menu.appendChild(btnTable);
    menu.appendChild(btnDelay);  // NEW: Add delay option
    menu.appendChild(btnCancel);
    
    document.body.appendChild(menu);
    __CONTEXT_MENU__ = menu;
    
    console.log('Extraction menu shown at', x, y);
  }
  
  function createMenuButton(text, bgColor) {
    const btn = document.createElement('button');
    btn.textContent = text;
    btn.style.cssText = `
      display: block; width: 100%; padding: 8px; margin: 4px 0;
      background: ${bgColor}; color: white; border: none; border-radius: 4px;
      cursor: pointer; font-size: 13px;
    `;
    return btn;
  }


  // Extract text content
  window.__extractText = function() {
    console.log('__extractText called');
    const element = __CONTEXT_TARGET__;
    console.log('Target element:', element);
    
    if (!element) {
      console.error('No element selected!');
      return;
    }
    
    removeExtractionMenu();
    
    // Show custom input dialog instead of prompt()
    showInputDialog("💾 Variable Name", "extracted_value", (varName) => {
      if (varName && varName.trim()) {
        const locators = getLocatorCandidates(element);
        
        // Get text - handle input/textarea differently
        let text = '';
        if (element.tagName === 'INPUT' || element.tagName === 'TEXTAREA' || element.tagName === 'SELECT') {
          text = (element.value || '').trim();
        } else {
          text = (element.innerText || element.textContent || '').trim();
        }
        
        // NEW: Show save options dialog
        showSaveOptionsDialog((saveConfig) => {
          console.log('Sending extract action:', {key: varName.trim(), text: text, saveConfig: saveConfig});
          
          send('extract', {
            url: location.href,
            key: varName.trim(),
            locators: locators,
            extract_type: 'text',
            sample_value: text.substring(0, 100),
            save_options: saveConfig  // NEW: Include save configuration
          });
          
          highlightExtracted(element, `✅ Extracting: "${text.substring(0, 30)}..."`);
        });
      } else {
        console.log('User cancelled or entered empty name');
      }
    });
  };

  // Extract attribute value
  window.__extractAttribute = function() {
    console.log('__extractAttribute called');
    const element = __CONTEXT_TARGET__;
    
    if (!element) {
      console.error('No element selected!');
      return;
    }
    
    removeExtractionMenu();
    
    // First ask for attribute name
    showInputDialog("🔍 Attribute Name", "href", (attr) => {
      if (attr && attr.trim()) {
        // Then ask for variable name
        showInputDialog("💾 Variable Name", attr.trim() + "_value", (varName) => {
          if (varName && varName.trim()) {
            const locators = getLocatorCandidates(element);
            const value = element.getAttribute(attr.trim()) || '';
            
            // NEW: Show save options dialog
            showSaveOptionsDialog((saveConfig) => {
              console.log('Sending extract attribute action:', {key: varName.trim(), attr: attr.trim(), value: value, saveConfig: saveConfig});
              
              send('extract', {
                url: location.href,
                key: varName.trim(),
                locators: locators,
                extract_type: 'attribute',
                attribute_name: attr.trim(),
                sample_value: value,
                save_options: saveConfig  // NEW: Include save configuration
              });
              
              highlightExtracted(element, `✅ Extracting ${attr}: "${value.substring(0, 30)}..."`);
            });
          }
        });
      }
    });
  };

  // Extract table data - Phase 8.3
  window.__extractTable = function() {
    console.log('🔍 __extractTable called');
    const element = __CONTEXT_TARGET__;
    console.log('🔍 Target element:', element);
    
    if (!element) {
      console.error('❌ No element selected!');
      alert('❌ No element selected! Please right-click on a table.');
      return;
    }
    
    removeExtractionMenu();
    console.log('🔍 Menu removed, detecting table...');
    
    // Detect table element
    const table = detectTableElement(element);
    console.log('🔍 Table detected:', table);
    
    if (!table) {
      alert('⚠️ No table found! Please right-click on or near a table element.');
      console.error('❌ No table found near element:', element);
      return;
    }
    
    console.log('✅ Table found, getting headers...');
    const headers = getTableHeaders(table);
    console.log('📊 Headers:', headers);
    
    // Show column selection dialog
    console.log('🔍 Calling showColumnSelectionDialog...');
    showColumnSelectionDialog(table, (config) => {
      console.log('📋 Dialog callback called with config:', config);
      
      if (!config) {
        console.log('User cancelled table extraction');
        return;
      }
      
      console.log('Table extraction config:', config);
      
      // Ask for variable name
      showInputDialog("💾 Variable Name", "table_data", (varName) => {
        if (varName && varName.trim()) {
          console.log('✅ Variable name entered:', varName);
          
          // Generate CSS selector for the table
          const tableSelector = generateStableCSS(table);
          console.log('✅ Table selector generated:', tableSelector);
          
          const locators = getLocatorCandidates(table);
          console.log('✅ Locators generated:', locators);
          
          const table_config = {
            table_selector: tableSelector,
            columns: config.columns,
            column_indices: config.column_indices,
            pagination: config.pagination
          };
          
          console.log('✅ Table config created:', table_config);
          
          // Show save options dialog
          showSaveOptionsDialog((saveConfig) => {
            console.log('💾 Save config:', saveConfig);
            console.log('📤 Sending extract_table action...');
            
            send('extract_table', {
              url: location.href,
              key: varName.trim(),
              locators: locators,
              table_config: table_config,
              save_options: saveConfig
            });
            
            highlightExtracted(table, `✅ Extracting table: ${config.columns.length} columns`);
          });
        } else {
          console.log('User cancelled or entered empty name');
        }
      });
    });
  };

  // Custom input dialog (works in automation unlike prompt())
  function showInputDialog(title, defaultValue, callback) {
    // Remove any existing dialog
    const existing = document.getElementById('__extraction_input__');
    if (existing) existing.remove();
    
    const overlay = document.createElement('div');
    overlay.id = '__extraction_input__';
    overlay.style.cssText = `
      position: fixed; top: 0; left: 0; right: 0; bottom: 0;
      background: rgba(0,0,0,0.5); z-index: 2147483647;
      display: flex; align-items: center; justify-content: center;
    `;
    
    const dialog = document.createElement('div');
    dialog.style.cssText = `
      background: white; padding: 20px; border-radius: 8px;
      box-shadow: 0 4px 12px rgba(0,0,0,0.3); min-width: 300px;
    `;
    
    dialog.innerHTML = `
      <div style="font-weight: bold; margin-bottom: 10px; font-size: 15px;">${title}</div>
      <input type="text" id="__extract_input_field__" value="${defaultValue}" 
             style="width: 100%; padding: 8px; border: 1px solid #ccc; border-radius: 4px; font-size: 14px; margin-bottom: 10px;">
      <div style="display: flex; justify-content: flex-end; gap: 8px;">
        <button id="__extract_cancel__" style="padding: 8px 16px; background: #f44336; color: white; border: none; border-radius: 4px; cursor: pointer;">Cancel</button>
        <button id="__extract_ok__" style="padding: 8px 16px; background: #4CAF50; color: white; border: none; border-radius: 4px; cursor: pointer;">OK</button>
      </div>
    `;
    
    overlay.appendChild(dialog);
    document.body.appendChild(overlay);
    
    const input = document.getElementById('__extract_input_field__');
    const okBtn = document.getElementById('__extract_ok__');
    const cancelBtn = document.getElementById('__extract_cancel__');
    
    // Focus and select input
    setTimeout(() => {
      input.focus();
      input.select();
    }, 100);
    
    const cleanup = () => {
      overlay.remove();
    };
    
    okBtn.onclick = (e) => {
      e.stopPropagation();
      const value = input.value;
      cleanup();
      callback(value);
    };
    
    cancelBtn.onclick = (e) => {
      e.stopPropagation();
      cleanup();
      callback(null);
    };
    
    input.onkeydown = (e) => {
      if (e.key === 'Enter') {
        e.stopPropagation();
        const value = input.value;
        cleanup();
        callback(value);
      } else if (e.key === 'Escape') {
        e.stopPropagation();
        cleanup();
        callback(null);
      }
    };
  }

  // NEW: Add Delay (Phase 9.1)
  window.__addDelay = function() {
    removeExtractionMenu();
    
    showInputDialog('⏱️ Add Delay (seconds)', '5', (value) => {
      if (!value) return;  // User cancelled
      
      // Validate: Must be number between 1-60
      const seconds = parseFloat(value);
      if (isNaN(seconds) || seconds < 1 || seconds > 60) {
        showMessage(`❌ Invalid delay: Must be 1-60 seconds`);
        return;
      }
      
      // Round to 1 decimal place
      const roundedSeconds = Math.round(seconds * 10) / 10;
      
      // Send wait action to Python using send() helper
      send('wait', {
        url: location.href,
        value: roundedSeconds.toString(),
        name: null  // Optional reason field for future
      });
      
      // Show confirmation
      showMessage(`⏱️ Added ${roundedSeconds}s delay`);
    });
  };

  // ============================================================================
  // PHASE 8.2: SAVE DIALOGS
  // ============================================================================
  
  // Show save options dialog (Step 1 of save flow)
  function showSaveOptionsDialog(callback) {
    removeExtractionMenu();
    
    const overlay = document.createElement('div');
    overlay.id = '__save_options_dialog__';
    overlay.style.cssText = `
      position: fixed; top: 0; left: 0; right: 0; bottom: 0;
      background: rgba(0,0,0,0.5); z-index: 2147483647;
      display: flex; align-items: center; justify-content: center;
    `;
    
    const dialog = document.createElement('div');
    dialog.style.cssText = `
      background: white; padding: 20px; border-radius: 8px;
      box-shadow: 0 4px 12px rgba(0,0,0,0.3); min-width: 300px;
    `;
    
    dialog.innerHTML = `
      <div style="font-weight: bold; margin-bottom: 15px; font-size: 16px;">💾 Save Extracted Data?</div>
      <div style="margin-bottom: 15px;">
        <label style="display: block; margin: 8px 0; cursor: pointer;">
          <input type="checkbox" id="save_excel" style="margin-right: 8px;">
          <span>Save to Excel (.xlsx)</span>
        </label>
        <label style="display: block; margin: 8px 0; cursor: pointer;">
          <input type="checkbox" id="save_word" style="margin-right: 8px;">
          <span>Save to Word (.docx)</span>
        </label>
        <label style="display: block; margin: 8px 0; cursor: pointer;">
          <input type="checkbox" id="save_txt" style="margin-right: 8px;">
          <span>Save to Text (.txt)</span>
        </label>
      </div>
      <div style="display: flex; justify-content: flex-end; gap: 8px;">
        <button id="__save_skip__" style="padding: 8px 16px; background: #999; color: white; border: none; border-radius: 4px; cursor: pointer;">Skip</button>
        <button id="__save_configure__" style="padding: 8px 16px; background: #4CAF50; color: white; border: none; border-radius: 4px; cursor: pointer;">Configure & Save</button>
      </div>
    `;
    
    overlay.appendChild(dialog);
    document.body.appendChild(overlay);
    
    // CRITICAL: Mark dialog so recorder ignores it
    overlay.setAttribute('data-webai-dialog', 'true');
    
    const skipBtn = document.getElementById('__save_skip__');
    const configBtn = document.getElementById('__save_configure__');
    
    const cleanup = () => overlay.remove();
    
    skipBtn.onclick = (e) => {
      e.stopPropagation();
      cleanup();
      callback(null);  // No save
    };
    
    configBtn.onclick = (e) => {
      e.stopPropagation();
      
      const formats = {
        excel: document.getElementById('save_excel').checked,
        word: document.getElementById('save_word').checked,
        txt: document.getElementById('save_txt').checked
      };
      
      // DEBUG: Log checkbox states
      console.log('📋 Checkbox values:', formats);
      console.log('  Excel checkbox:', document.getElementById('save_excel'), 'checked:', document.getElementById('save_excel')?.checked);
      console.log('  Word checkbox:', document.getElementById('save_word'), 'checked:', document.getElementById('save_word')?.checked);
      console.log('  Txt checkbox:', document.getElementById('save_txt'), 'checked:', document.getElementById('save_txt')?.checked);
      
      if (!formats.excel && !formats.word && !formats.txt) {
        cleanup();
        callback(null);
        return;
      }
      
      cleanup();
      showFileConfigDialog(formats, callback);
    };
  }
  
  // Show file configuration dialog (Step 2 of save flow)
  function showFileConfigDialog(formats, callback) {
    const overlay = document.createElement('div');
    overlay.id = '__file_config_dialog__';
    overlay.style.cssText = `
      position: fixed; top: 0; left: 0; right: 0; bottom: 0;
      background: rgba(0,0,0,0.5); z-index: 2147483647;
      display: flex; align-items: center; justify-content: center;
    `;
    
    const dialog = document.createElement('div');
    dialog.style.cssText = `
      background: white; padding: 20px; border-radius: 8px;
      box-shadow: 0 4px 12px rgba(0,0,0,0.3); min-width: 400px;
    `;
    
    const currentDir = 'e:\\\\webai_playwright_python';
    
    dialog.innerHTML = `
      <div style="font-weight: bold; margin-bottom: 15px; font-size: 16px;">📁 File Configuration</div>
      <div style="margin-bottom: 15px;">
        <label style="display: block; margin-bottom: 5px; font-weight: 500;">Folder Path:</label>
        <input type="text" id="__file_folder__" value="${currentDir}" 
               style="width: 100%; padding: 8px; border: 1px solid #ccc; border-radius: 4px; font-size: 14px;">
        <div style="font-size: 12px; color: #666; margin-top: 3px;">Will be created if it doesn't exist</div>
      </div>
      <div style="margin-bottom: 15px;">
        <label style="display: block; margin-bottom: 5px; font-weight: 500;">Filename (without extension):</label>
        <input type="text" id="__file_name__" value="extracted_data" 
               style="width: 100%; padding: 8px; border: 1px solid #ccc; border-radius: 4px; font-size: 14px;">
      </div>
      <div style="margin-bottom: 15px;">
        <div style="font-weight: 500; margin-bottom: 8px;">Mode:</div>
        <label style="display: block; margin: 5px 0; cursor: pointer;">
          <input type="radio" name="__file_mode__" value="append" checked style="margin-right: 8px;">
          <span>Append to existing file</span>
        </label>
        <label style="display: block; margin: 5px 0; cursor: pointer;">
          <input type="radio" name="__file_mode__" value="new" style="margin-right: 8px;">
          <span>Create new file (overwrite)</span>
        </label>
      </div>
      <div style="display: flex; justify-content: flex-end; gap: 8px;">
        <button id="__file_cancel__" style="padding: 8px 16px; background: #f44336; color: white; border: none; border-radius: 4px; cursor: pointer;">Cancel</button>
        <button id="__file_save__" style="padding: 8px 16px; background: #4CAF50; color: white; border: none; border-radius: 4px; cursor: pointer;">Save</button>
      </div>
    `;
    
    overlay.appendChild(dialog);
    document.body.appendChild(overlay);
    
    // CRITICAL: Mark dialog so recorder ignores it
    overlay.setAttribute('data-webai-dialog', 'true');
    
    const folderInput = document.getElementById('__file_folder__');
    const nameInput = document.getElementById('__file_name__');
    const cancelBtn = document.getElementById('__file_cancel__');
    const saveBtn = document.getElementById('__file_save__');
    
    const cleanup = () => overlay.remove();
    
    cancelBtn.onclick = (e) => {
      e.stopPropagation();
      cleanup();
      callback(null);
    };
    
    saveBtn.onclick = (e) => {
      e.stopPropagation();
      
      const folder = folderInput.value.trim();
      const filename = nameInput.value.trim();
      const mode = document.querySelector('input[name="__file_mode__"]:checked').value;
      
      if (!folder || !filename) {
        alert('Please enter both folder and filename');
        return;
      }
      
      cleanup();
      callback({
        formats: formats,
        folder: folder,
        filename: filename,
        mode: mode
      });
    };
  }

  // Visual feedback for extraction
  function highlightExtracted(element, message) {
    const originalOutline = element.style.outline;
    const originalBg = element.style.backgroundColor;
    
    element.style.outline = '3px solid #4CAF50';
    element.style.backgroundColor = '#E8F5E9';
    
    // Show message
    showHint(message);
    
    setTimeout(() => {
      element.style.outline = originalOutline;
      element.style.backgroundColor = originalBg;
    }, 2000);
  }

  function removeExtractionMenu() {
    if (__CONTEXT_MENU__) {
      __CONTEXT_MENU__.remove();
      __CONTEXT_MENU__ = null;
    }
  }

  window.__cancelExtract = removeExtractionMenu;

  // Close menu on any click outside
  document.addEventListener('click', (e) => {
    if (__CONTEXT_MENU__ && !__CONTEXT_MENU__.contains(e.target)) {
      removeExtractionMenu();
    }
  }, true);

})();
"""


class WebRecorder:
    def __init__(self) -> None:
        self.steps: List[Step] = []
        self._stop_event = asyncio.Event()
        self._last_url: Optional[str] = None
        self._opened_once = False

    async def attach(self, page: Page) -> None:
        async def on_event(source, payload: Dict[str, Any]) -> None:
            kind = payload.get("kind")
            url = payload.get("url")

            if url and url != self._last_url:
                self._last_url = url
                action = "open" if not self._opened_once else "navigate"
                self._opened_once = True
                self.steps.append(Step(action=action, url=url, ts=time.time()))

            if kind == "click":
                name = payload.get("name")
                locators = payload.get("locators")  # NEW: Extract locators
                if name:
                    if self.steps and self.steps[-1].action == "click" and self.steps[-1].url == url and self.steps[-1].name == name:
                        return
                    self.steps.append(Step(action="click", url=url, name=name, locators=locators, ts=time.time()))

            elif kind == "type_final":
                label = payload.get("label")
                value = payload.get("value")
                locators = payload.get("locators")  # NEW: Extract locators
                # coalesce repeated types on same url+label
                if self.steps and self.steps[-1].action == "type" and self.steps[-1].url == url and self.steps[-1].name == label:
                    self.steps[-1].value = value
                    self.steps[-1].locators = locators  # NEW: Update locators
                    self.steps[-1].ts = time.time()
                    return
                self.steps.append(Step(action="type", url=url, name=label, value=value, locators=locators, ts=time.time()))

            elif kind == "extract":
                key = payload.get("key")
                extract_type = payload.get("extract_type")
                locators = payload.get("locators")
                sample_value = payload.get("sample_value")
                attribute_name = payload.get("attribute_name")
                save_options = payload.get("save_options")  # NEW: Get save config
                
                # Create a Step for extraction
                step = Step(
                    action="extract",
                    url=url,
                    name=key,  # Variable name
                    value=sample_value,  # Sample of extracted value
                    locators=locators,
                    extract_type=extract_type,  # 'text' or 'attribute'
                    attribute_name=attribute_name,  # Only set for attribute extraction
                    save_options=save_options,  # NEW: Include save configuration
                    ts=time.time()
                )
                
                self.steps.append(step)
                print(f"✅ Recorded extraction: {key} = '{sample_value[:50] if sample_value else 'N/A'}'")
                
                # NEW: Phase 8.2 - Handle immediate save if configured
                if save_options:
                    self._save_extraction_immediately(step, save_options)
            
            elif kind == "extract_table":
                # NEW: Phase 8.3 - Table extraction
                key = payload.get("key")
                locators = payload.get("locators")
                table_config = payload.get("table_config")
                save_options = payload.get("save_options")
                
                # Create a Step for table extraction
                step = Step(
                    action="extract_table",
                    url=url,
                    name=key,  # Variable name
                    locators=locators,
                    table_config=table_config,  # Table configuration
                    save_options=save_options,  # Save configuration
                    ts=time.time()
                )
                
                self.steps.append(step)
                cols_count = len(table_config.get('columns', [])) if table_config else 0
                print(f"✅ Recorded table extraction: {key} ({cols_count} columns)")
            
            elif kind == "wait":
                # NEW: Handle explicit wait/delay
                seconds = payload.get("value")
                name = payload.get("name")  # Optional reason for wait
                
                try:
                    wait_time = float(seconds)
                    if wait_time < 1 or wait_time > 60:
                        print(f"⚠️ Invalid wait time: {wait_time}s (must be 1-60)")
                        return
                    
                    step = Step(
                        action="wait",
                        url=url,
                        name=name,  # Optional description
                        value=str(wait_time),  # Wait duration in seconds
                        ts=time.time()
                    )
                    
                    self.steps.append(step)
                    print(f"⏱️ Recorded delay: {wait_time}s")
                except (ValueError, TypeError):
                    print(f"⚠️ Invalid wait value: {seconds}")

        async def stop(source, payload=None) -> None:
            self._stop_event.set()

        async def verify_text(source, payload: Dict[str, Any]) -> None:
            self.steps.append(Step(action="verify_text", url=payload.get("url"), value=payload.get("text"), ts=time.time()))

        async def verify_visible(source, payload: Dict[str, Any]) -> None:
            self.steps.append(Step(action="verify_visible", url=payload.get("url"), value=payload.get("text"), ts=time.time()))

        await page.expose_binding("__recordEvent", on_event)
        await page.expose_binding("__stopRecording", stop)
        await page.expose_binding("__recordVerifyText", verify_text)
        await page.expose_binding("__recordVerifyVisible", verify_visible)

        await page.add_init_script(RECORDER_INIT_SCRIPT)

    async def wait_for_stop(self) -> None:
        await self._stop_event.wait()

    def to_json(self) -> List[Dict[str, Any]]:
        return [asdict(s) for s in self.steps]

    def to_task_text(self) -> str:
        lines: List[str] = []
        for s in self.steps:
            if s.action == "open" and s.url:
                lines.append(f"Open {s.url}")
            elif s.action == "navigate" and s.url:
                lines.append(f"Navigate to {s.url}")
            elif s.action == "click" and s.name:
                lines.append(f'Click "{s.name}"')
            elif s.action == "type":
                if s.name:
                    lines.append(f'Type "{s.value}" into "{s.name}"')
                else:
                    lines.append(f'Type "{s.value}"')
            elif s.action == "press_key":
                # Handle key press actions
                key = s.key or "Enter"  # fallback to Enter if key not specified
                context = s.name
                if context:
                    lines.append(f'Press {key} in "{context}"')
                else:
                    lines.append(f'Press {key}')
            elif s.action == "verify_text":
                lines.append(f'Verify the page contains text "{s.value}"')
            elif s.action == "verify_visible":
                lines.append(f'Verify element "{s.value}" is visible')
            elif s.action == "extract":
                # NEW: Handle extraction actions with element context
                element_context = ""
                
                # Try to get element context from locators
                if s.locators:
                    # Prioritize text or aria-label for context
                    for loc in s.locators:
                        if loc.get("type") == "text" and loc.get("value"):
                            element_context = f" from '{loc['value']}' element"
                            break
                        elif loc.get("type") == "aria-label" and loc.get("value"):
                            element_context = f" from '{loc['value']}' element"
                            break
                        elif loc.get("type") == "role" and loc.get("name"):
                            element_context = f" from '{loc['name']}' {loc['value']}"
                            break
                
                if s.extract_type == "attribute" and s.attribute_name:
                    lines.append(f'Extract attribute "{s.attribute_name}"{element_context} into variable "{s.name}"')
                else:
                    lines.append(f'Extract text{element_context} into variable "{s.name}"')
            elif s.action == "wait":
                # NEW: Handle wait/delay actions
                seconds = s.value
                if s.name:  # Optional reason/description
                    lines.append(f'Wait {seconds} seconds ({s.name})')
                else:
                    lines.append(f'Wait {seconds} seconds')

        lines.append("Finish with done only after verification.")

        seen = set()
        out = []
        for ln in lines:
            if ln not in seen:
                seen.add(ln)
                out.append(ln)
        return "\n".join(out)
    
    # ============================================================================
    # PHASE 8.2: SAVE EXTRACTED DATA
    # ============================================================================
    
    def _save_extraction_immediately(self, step: Step, options: Dict[str, Any]) -> None:
        """Save single extraction to files immediately"""
        import os
        from datetime import datetime
        
        formats = options.get('formats', {})
        folder = options.get('folder', '.')
        filename = options.get('filename', 'extracted_data')
        mode = options.get('mode', 'append')
        
        # Auto-create folder if it doesn't exist
        if not os.path.exists(folder):
            try:
                os.makedirs(folder, exist_ok=True)
                print(f"✅ Created folder: {folder}")
            except Exception as e:
                print(f"⚠️ Cannot create folder ({e}), using current directory")
                folder = '.'
        
        saved_files = []
        
        try:
            if formats.get('excel'):
                filepath = os.path.join(folder, f"{filename}.xlsx")
                self._save_to_excel_immediate(step, filepath, mode)
                saved_files.append(os.path.basename(filepath))
        except Exception as e:
            print(f"⚠️ Excel save failed: {e}")
        
        try:
            if formats.get('word'):
                filepath = os.path.join(folder, f"{filename}.docx")
                self._save_to_word_immediate(step, filepath, mode)
                saved_files.append(os.path.basename(filepath))
        except Exception as e:
            print(f"⚠️ Word save failed: {e}")
        
        try:
            if formats.get('txt'):
                filepath = os.path.join(folder, f"{filename}.txt")
                self._save_to_txt_immediate(step, filepath, mode)
                saved_files.append(os.path.basename(filepath))
        except Exception as e:
            print(f"⚠️ Text save failed: {e}")
        
        if saved_files:
            print(f"💾 Saved to: {', '.join(saved_files)}")
    
    def _save_to_excel_immediate(self, step: Step, filepath: str, mode: str) -> None:
        """Save extraction to Excel file"""
        try:
            import openpyxl
            from openpyxl.styles import Font
        except ImportError:
            print("⚠️ openpyxl not installed. Run: pip install openpyxl")
            return
        
        import os
        from datetime import datetime
        
        if mode == 'new' or not os.path.exists(filepath):
            # Create new workbook
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Extracted Data"
            
            # Add headers
            headers = ['Variable', 'Value', 'Type', 'Attribute', 'URL', 'Timestamp']
            ws.append(headers)
            
            # Format headers
            for cell in ws[1]:
                cell.font = Font(bold=True)
        else:
            # Append to existing
            wb = openpyxl.load_workbook(filepath)
            ws = wb.active
        
        # Add data row
        ws.append([
            step.name,
            step.value or '',
            step.extract_type or 'text',
            step.attribute_name or '',
            step.url or '',
            datetime.fromtimestamp(step.ts).strftime('%Y-%m-%d %H:%M:%S')
        ])
        
        # Auto-adjust column widths
        for column in ws.columns:
            max_length = 0
            column = [cell for cell in column]
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(cell.value)
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)
            ws.column_dimensions[column[0].column_letter].width = adjusted_width
        
        wb.save(filepath)
    
    def _save_to_word_immediate(self, step: Step, filepath: str, mode: str) -> None:
        """Save extraction to Word file"""
        try:
            from docx import Document
            from docx.shared import Pt
        except ImportError:
            print("⚠️ python-docx not installed. Run: pip install python-docx")
            return
        
        import os
        from datetime import datetime
        
        if mode == 'new' or not os.path.exists(filepath):
            # Create new document
            doc = Document()
            doc.add_heading('Extracted Data', 0)
        else:
            # Load existing
            doc = Document(filepath)
        
        # Add extraction entry
        doc.add_paragraph('━' * 60)
        
        heading = doc.add_heading(f"Variable: {step.name}", level=2)
        
        doc.add_paragraph(f"Value: {step.value or '(empty)'}")
        doc.add_paragraph(f"Type: {step.extract_type or 'text'}")
        
        if step.attribute_name:
            doc.add_paragraph(f"Attribute: {step.attribute_name}")
        
        doc.add_paragraph(f"URL: {step.url or ''}")
        doc.add_paragraph(f"Timestamp: {datetime.fromtimestamp(step.ts).strftime('%Y-%m-%d %H:%M:%S')}")
        
        doc.save(filepath)
    
    def _save_to_txt_immediate(self, step: Step, filepath: str, mode: str) -> None:
        """Save extraction to text file"""
        import os
        from datetime import datetime
        
        write_mode = 'w' if mode == 'new' and os.path.exists(filepath) else 'a'
        
        with open(filepath, write_mode, encoding='utf-8') as f:
            if write_mode == 'w' or not os.path.exists(filepath) or os.path.getsize(filepath) == 0:
                f.write(f"=== Extracted Data Log ===\n")
                f.write(f"Started: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
            
            f.write(f"{'─' * 60}\n")
            f.write(f"Variable: {step.name}\n")
            f.write(f"Value: {step.value or '(empty)'}\n")
            f.write(f"Type: {step.extract_type or 'text'}\n")
            
            if step.attribute_name:
                f.write(f"Attribute: {step.attribute_name}\n")
            
            f.write(f"URL: {step.url or ''}\n")
            f.write(f"Time: {datetime.fromtimestamp(step.ts).strftime('%Y-%m-%d %H:%M:%S')}\n\n")
