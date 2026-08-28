"""
Data Extraction Plugin for WebAI Playwright Recorder.

Handles data extraction UI (context menu, text/attribute/table extraction dialogs)
and immediate persistence of extracted data into Excel (.xlsx), Word (.docx), or Text (.txt) files.
"""
from __future__ import annotations

import os
from datetime import datetime
from typing import Any, Dict, Optional, TYPE_CHECKING

if TYPE_CHECKING:
    from ..recorder import Step, WebRecorder

EXTRACTION_INIT_SCRIPT = r"""
(() => {
  // ============================================================================
  // TABLE EXTRACTION SYSTEM - Phase 8.3
  // ============================================================================
  
  function detectTableElement(target) {
    let el = target;
    while (el && el !== document.body) {
      if (el.tagName === 'TABLE') return el;
      if (el.querySelector('table')) return el.querySelector('table');
      if (el.getAttribute && el.getAttribute('role') === 'table') return el;
      el = el.parentElement;
    }
    return null;
  }
  
  function getTableHeaders(table) {
    let headers = Array.from(table.querySelectorAll('thead th, thead td'))
      .map(h => h.innerText.trim());
    
    if (headers.length === 0) {
      const firstRow = table.querySelector('tr');
      if (firstRow) {
        headers = Array.from(firstRow.querySelectorAll('th, td'))
          .map(h => h.innerText.trim());
      }
    }
    
    return headers.filter(h => h.length > 0);
  }
  
  function showColumnSelectionDialog(table, callback) {
    const headers = getTableHeaders(table);
    
    if (headers.length === 0) {
      alert('No table headers found!');
      callback(null);
      return;
    }
    
    const overlay = document.createElement('div');
    overlay.style.cssText = `
      position: fixed; top: 0; left: 0; right: 0; bottom: 0;
      background: rgba(0,0,0,0.5); z-index: 999999;
      display: flex; align-items: center; justify-content: center;
    `;
    overlay.setAttribute('data-webai-dialog', 'true');
    
    const dialog = document.createElement('div');
    dialog.style.cssText = `
      background: white; padding: 20px; border-radius: 8px;
      box-shadow: 0 4px 12px rgba(0,0,0,0.3);
      min-width: 400px; max-width: 600px; max-height: 80vh;
      overflow-y: auto;
    `;
    
    dialog.innerHTML = `
      <div style="font-weight: bold; margin-bottom: 15px; font-size: 16px;">
        Select Columns to Extract
      </div>
      <div style="margin-bottom: 10px;">
        <button id="__select_all__" style="margin-right: 8px; padding: 4px 8px; cursor: pointer;">Select All</button>
        <button id="__deselect_all__" style="padding: 4px 8px; cursor: pointer;">Deselect All</button>
      </div>
      <div id="__column_list__" style="margin-bottom: 15px; max-height: 300px; overflow-y: auto; border: 1px solid #ddd; padding: 10px; border-radius: 4px;">
        ${headers.map((h, i) => `
          <label style="display: block; margin: 8px 0; cursor: pointer;">
            <input type="checkbox" class="column-checkbox" data-index="${i}" 
                   value="${h}" style="margin-right: 8px;" checked>
            <span>${h || ('Column ' + (i + 1))}</span>
          </label>
        `).join('')}
      </div>
      <div style="margin-bottom: 15px; padding: 10px; background: #f5f5f5; border-radius: 4px;">
        <label style="cursor: pointer; display: flex; align-items: center; margin-bottom: 8px;">
          <input type="checkbox" id="__pagination_enabled__" style="margin-right: 8px;">
          <span style="font-weight: bold;">Enable Pagination</span>
        </label>
        <div id="__pagination_options__" style="display: none; margin-left: 24px;">
          <label style="display: block; margin: 4px 0;">
            <span>Max Pages (1-100):</span>
            <input type="number" id="__max_pages__" value="10" min="1" max="100" 
                   style="width: 60px; margin-left: 8px; padding: 2px;">
          </label>
          <label style="display: block; margin: 4px 0;">
            <span>Wait Per Page (1-10 sec):</span>
            <input type="number" id="__wait_per_page__" value="2" min="1" max="10" step="0.5"
                   style="width: 60px; margin-left: 8px; padding: 2px;">
          </label>
          <label style="display: block; margin: 4px 0;">
            <span>Page Timeout (5-30 sec):</span>
            <input type="number" id="__page_timeout__" value="10" min="5" max="30" step="1"
                   style="width: 60px; margin-left: 8px; padding: 2px;">
          </label>
          <label style="display: block; margin: 4px 0;">
            <span>Retry Attempts (1-5):</span>
            <input type="number" id="__retry_attempts__" value="3" min="1" max="5" step="1"
                   style="width: 60px; margin-left: 8px; padding: 2px;">
          </label>
        </div>
      </div>
      <div style="display: flex; justify-content: flex-end; gap: 8px;">
        <button id="__column_cancel__" style="padding: 8px 16px; background: #999; color: white; border: none; border-radius: 4px; cursor: pointer;">Cancel</button>
        <button id="__column_extract__" style="padding: 8px 16px; background: #4CAF50; color: white; border: none; border-radius: 4px; cursor: pointer;">Extract</button>
      </div>
    `;
    
    overlay.appendChild(dialog);
    document.body.appendChild(overlay);
    
    const cleanup = () => {
      if (overlay.parentElement) {
        overlay.parentElement.removeChild(overlay);
      }
    };
    
    document.getElementById('__select_all__').onclick = () => {
      document.querySelectorAll('.column-checkbox').forEach(cb => cb.checked = true);
    };
    
    document.getElementById('__deselect_all__').onclick = () => {
      document.querySelectorAll('.column-checkbox').forEach(cb => cb.checked = false);
    };
    
    const paginationCheckbox = document.getElementById('__pagination_enabled__');
    const paginationOptions = document.getElementById('__pagination_options__');
    paginationCheckbox.onchange = () => {
      paginationOptions.style.display = paginationCheckbox.checked ? 'block' : 'none';
    };
    
    document.getElementById('__column_cancel__').onclick = () => {
      cleanup();
      callback(null);
    };
    
    document.getElementById('__column_extract__').onclick = () => {
      const selectedColumns = [];
      const selectedIndices = [];
      
      document.querySelectorAll('.column-checkbox:checked').forEach(cb => {
        selectedColumns.push(cb.value);
        selectedIndices.push(parseInt(cb.getAttribute('data-index')));
      });
      
      if (selectedColumns.length === 0) {
        alert('Please select at least one column!');
        return;
      }
      
      const config = {
        columns: selectedColumns,
        column_indices: selectedIndices,
        pagination: {
          enabled: paginationCheckbox.checked,
          max_pages: parseInt(document.getElementById('__max_pages__').value) || 10,
          wait_per_page: parseFloat(document.getElementById('__wait_per_page__').value) || 2.0,
          page_timeout: parseFloat(document.getElementById('__page_timeout__').value) || 10.0,
          retry_attempts: parseInt(document.getElementById('__retry_attempts__').value) || 3
        }
      };
      
      cleanup();
      callback(config);
    };
  }

  // ============================================================================
  // DATA EXTRACTION SYSTEM - Right-Click Context Menu
  // ============================================================================
  
  let __CONTEXT_MENU__ = null;
  let __CONTEXT_TARGET__ = null;

  document.addEventListener('contextmenu', (e) => {
    if (window.__RECORDER_STOPPED__) return;
    
    e.preventDefault();
    e.stopPropagation();
    
    __CONTEXT_TARGET__ = e.target;
    showExtractionMenu(e.pageX, e.pageY);
  }, true);

  function showExtractionMenu(x, y) {
    removeExtractionMenu();
    
    const menu = document.createElement('div');
    menu.id = '__extraction_menu__';
    menu.style.cssText = `
      position: fixed;
      left: ${x}px;
      top: ${y}px;
      z-index: 2147483647;
      background: white;
      border: 2px solid #4CAF50;
      border-radius: 8px;
      padding: 12px;
      box-shadow: 0 4px 12px rgba(0,0,0,0.3);
      font-family: Arial, sans-serif;
      font-size: 14px;
    `;
    
    const title = document.createElement('div');
    title.style.cssText = 'font-weight: bold; margin-bottom: 10px; color: #333; font-size: 15px;';
    title.textContent = 'Extract Data';
    
    const btnText = createMenuButton('Extract Text', '#4CAF50');
    btnText.addEventListener('click', (e) => {
      e.stopPropagation();
      console.log('Extract Text clicked');
      window.__extractText();
    });
    
    const btnAttr = createMenuButton('Extract Attribute', '#2196F3');
    btnAttr.addEventListener('click', (e) => {
      e.stopPropagation();
      console.log('Extract Attribute clicked');
      window.__extractAttribute();
    });
    
    const btnTable = createMenuButton('Extract Table', '#FF9800');
    btnTable.addEventListener('click', (e) => {
      e.stopPropagation();
      console.log('Extract Table clicked');
      window.__extractTable();
    });
    
    const btnDelay = createMenuButton('Add Delay', '#9C27B0');
    btnDelay.addEventListener('click', (e) => {
      e.stopPropagation();
      console.log('Add Delay clicked');
      window.__addDelay();
    });
    
    const btnCancel = createMenuButton('Cancel', '#f44336');
    btnCancel.addEventListener('click', (e) => {
      e.stopPropagation();
      console.log('Cancel clicked');
      window.__cancelExtract();
    });
    
    menu.appendChild(title);
    menu.appendChild(btnText);
    menu.appendChild(btnAttr);
    menu.appendChild(btnTable);
    menu.appendChild(btnDelay);
    menu.appendChild(btnCancel);
    
    document.body.appendChild(menu);
    __CONTEXT_MENU__ = menu;
  }
  
  function createMenuButton(text, bgColor) {
    const btn = document.createElement('button');
    btn.textContent = text;
    btn.style.cssText = `
      display: block; width: 100%; padding: 8px; margin: 4px 0;
      background: ${bgColor}; color: white; border: none; border-radius: 4px;
      cursor: pointer; font-size: 13px;
    `;
    return btn;
  }

  // Extract text content
  window.__extractText = function() {
    const element = __CONTEXT_TARGET__;
    if (!element) return;
    
    removeExtractionMenu();
    
    showInputDialog("Variable Name", "extracted_value", (varName) => {
      if (varName && varName.trim()) {
        const locators = typeof window.getLocatorCandidates === "function" 
          ? window.getLocatorCandidates(element) 
          : [];
        
        let text = '';
        if (element.tagName === 'INPUT' || element.tagName === 'TEXTAREA' || element.tagName === 'SELECT') {
          text = (element.value || '').trim();
        } else {
          text = (element.innerText || element.textContent || '').trim();
        }
        
        showSaveOptionsDialog((saveConfig) => {
          if (typeof window.__recordSend === "function") {
            window.__recordSend('extract', {
              url: location.href,
              key: varName.trim(),
              locators: locators,
              extract_type: 'text',
              sample_value: text.substring(0, 100),
              save_options: saveConfig
            });
          }
          
          highlightExtracted(element, `Extracting: "${text.substring(0, 30)}..."`);
        });
      }
    });
  };

  // Extract attribute value
  window.__extractAttribute = function() {
    const element = __CONTEXT_TARGET__;
    if (!element) return;
    
    removeExtractionMenu();
    
    showInputDialog("Attribute Name", "href", (attr) => {
      if (attr && attr.trim()) {
        showInputDialog("Variable Name", attr.trim() + "_value", (varName) => {
          if (varName && varName.trim()) {
            const locators = typeof window.getLocatorCandidates === "function" 
              ? window.getLocatorCandidates(element) 
              : [];
            const value = element.getAttribute(attr.trim()) || '';
            
            showSaveOptionsDialog((saveConfig) => {
              if (typeof window.__recordSend === "function") {
                window.__recordSend('extract', {
                  url: location.href,
                  key: varName.trim(),
                  locators: locators,
                  extract_type: 'attribute',
                  attribute_name: attr.trim(),
                  sample_value: value,
                  save_options: saveConfig
                });
              }
              
              highlightExtracted(element, `Extracting ${attr}: "${value.substring(0, 30)}..."`);
            });
          }
        });
      }
    });
  };

  // Extract table data - Phase 8.3
  window.__extractTable = function() {
    const element = __CONTEXT_TARGET__;
    if (!element) {
      alert('No element selected! Please right-click on a table.');
      return;
    }
    
    removeExtractionMenu();
    
    const table = detectTableElement(element);
    if (!table) {
      alert('No table found! Please right-click on or near a table element.');
      return;
    }
    
    showColumnSelectionDialog(table, (config) => {
      if (!config) return;
      
      showInputDialog("Variable Name", "table_data", (varName) => {
        if (varName && varName.trim()) {
          const tableSelector = typeof window.generateStableCSS === "function"
            ? window.generateStableCSS(table)
            : null;
          
          const locators = typeof window.getLocatorCandidates === "function"
            ? window.getLocatorCandidates(table)
            : [];
          
          const table_config = {
            table_selector: tableSelector,
            columns: config.columns,
            column_indices: config.column_indices,
            pagination: config.pagination
          };
          
          showSaveOptionsDialog((saveConfig) => {
            if (typeof window.__recordSend === "function") {
              window.__recordSend('extract_table', {
                url: location.href,
                key: varName.trim(),
                locators: locators,
                table_config: table_config,
                save_options: saveConfig
              });
            }
            
            highlightExtracted(table, `Extracting table: ${config.columns.length} columns`);
          });
        }
      });
    });
  };

  function showInputDialog(title, defaultValue, callback) {
    const existing = document.getElementById('__extraction_input__');
    if (existing) existing.remove();
    
    const overlay = document.createElement('div');
    overlay.id = '__extraction_input__';
    overlay.style.cssText = `
      position: fixed; top: 0; left: 0; right: 0; bottom: 0;
      background: rgba(0,0,0,0.5); z-index: 2147483647;
      display: flex; align-items: center; justify-content: center;
    `;
    
    const dialog = document.createElement('div');
    dialog.style.cssText = `
      background: white; padding: 20px; border-radius: 8px;
      box-shadow: 0 4px 12px rgba(0,0,0,0.3); min-width: 300px;
    `;
    
    dialog.innerHTML = `
      <div style="font-weight: bold; margin-bottom: 10px; font-size: 15px;">${title}</div>
      <input type="text" id="__extract_input_field__" value="${defaultValue}" 
             style="width: 100%; padding: 8px; border: 1px solid #ccc; border-radius: 4px; font-size: 14px; margin-bottom: 10px;">
      <div style="display: flex; justify-content: flex-end; gap: 8px;">
        <button id="__extract_cancel__" style="padding: 8px 16px; background: #f44336; color: white; border: none; border-radius: 4px; cursor: pointer;">Cancel</button>
        <button id="__extract_ok__" style="padding: 8px 16px; background: #4CAF50; color: white; border: none; border-radius: 4px; cursor: pointer;">OK</button>
      </div>
    `;
    
    overlay.appendChild(dialog);
    document.body.appendChild(overlay);
    
    const input = document.getElementById('__extract_input_field__');
    const okBtn = document.getElementById('__extract_ok__');
    const cancelBtn = document.getElementById('__extract_cancel__');
    
    setTimeout(() => {
      input.focus();
      input.select();
    }, 100);
    
    const cleanup = () => overlay.remove();
    
    okBtn.onclick = (e) => {
      e.stopPropagation();
      const value = input.value;
      cleanup();
      callback(value);
    };
    
    cancelBtn.onclick = (e) => {
      e.stopPropagation();
      cleanup();
      callback(null);
    };
    
    input.onkeydown = (e) => {
      if (e.key === 'Enter') {
        e.stopPropagation();
        const value = input.value;
        cleanup();
        callback(value);
      } else if (e.key === 'Escape') {
        e.stopPropagation();
        cleanup();
        callback(null);
      }
    };
  }

  // ============================================================================
  // SAVE DIALOGS
  // ============================================================================
  
  function showSaveOptionsDialog(callback) {
    removeExtractionMenu();
    
    const overlay = document.createElement('div');
    overlay.id = '__save_options_dialog__';
    overlay.style.cssText = `
      position: fixed; top: 0; left: 0; right: 0; bottom: 0;
      background: rgba(0,0,0,0.5); z-index: 2147483647;
      display: flex; align-items: center; justify-content: center;
    `;
    
    const dialog = document.createElement('div');
    dialog.style.cssText = `
      background: white; padding: 20px; border-radius: 8px;
      box-shadow: 0 4px 12px rgba(0,0,0,0.3); min-width: 300px;
    `;
    
    dialog.innerHTML = `
      <div style="font-weight: bold; margin-bottom: 15px; font-size: 16px;">Save Extracted Data?</div>
      <div style="margin-bottom: 15px;">
        <label style="display: block; margin: 8px 0; cursor: pointer;">
          <input type="checkbox" id="save_excel" style="margin-right: 8px;">
          <span>Save to Excel (.xlsx)</span>
        </label>
        <label style="display: block; margin: 8px 0; cursor: pointer;">
          <input type="checkbox" id="save_word" style="margin-right: 8px;">
          <span>Save to Word (.docx)</span>
        </label>
        <label style="display: block; margin: 8px 0; cursor: pointer;">
          <input type="checkbox" id="save_txt" style="margin-right: 8px;">
          <span>Save to Text (.txt)</span>
        </label>
      </div>
      <div style="display: flex; justify-content: flex-end; gap: 8px;">
        <button id="__save_skip__" style="padding: 8px 16px; background: #999; color: white; border: none; border-radius: 4px; cursor: pointer;">Skip</button>
        <button id="__save_configure__" style="padding: 8px 16px; background: #4CAF50; color: white; border: none; border-radius: 4px; cursor: pointer;">Configure & Save</button>
      </div>
    `;
    
    overlay.appendChild(dialog);
    document.body.appendChild(overlay);
    overlay.setAttribute('data-webai-dialog', 'true');
    
    const skipBtn = document.getElementById('__save_skip__');
    const configBtn = document.getElementById('__save_configure__');
    
    const cleanup = () => overlay.remove();
    
    skipBtn.onclick = (e) => {
      e.stopPropagation();
      cleanup();
      callback(null);
    };
    
    configBtn.onclick = (e) => {
      e.stopPropagation();
      
      const formats = {
        excel: document.getElementById('save_excel').checked,
        word: document.getElementById('save_word').checked,
        txt: document.getElementById('save_txt').checked
      };
      
      if (!formats.excel && !formats.word && !formats.txt) {
        cleanup();
        callback(null);
        return;
      }
      
      cleanup();
      showFileConfigDialog(formats, callback);
    };
  }
  
  function showFileConfigDialog(formats, callback) {
    const overlay = document.createElement('div');
    overlay.id = '__file_config_dialog__';
    overlay.style.cssText = `
      position: fixed; top: 0; left: 0; right: 0; bottom: 0;
      background: rgba(0,0,0,0.5); z-index: 2147483647;
      display: flex; align-items: center; justify-content: center;
    `;
    
    const dialog = document.createElement('div');
    dialog.style.cssText = `
      background: white; padding: 20px; border-radius: 8px;
      box-shadow: 0 4px 12px rgba(0,0,0,0.3); min-width: 400px;
    `;
    
    const currentDir = '.';
    
    dialog.innerHTML = `
      <div style="font-weight: bold; margin-bottom: 15px; font-size: 16px;">📁 File Configuration</div>
      <div style="margin-bottom: 15px;">
        <label style="display: block; margin-bottom: 5px; font-weight: 500;">Folder Path:</label>
        <input type="text" id="__file_folder__" value="${currentDir}" 
               style="width: 100%; padding: 8px; border: 1px solid #ccc; border-radius: 4px; font-size: 14px;">
        <div style="font-size: 12px; color: #666; margin-top: 3px;">Will be created if it doesn't exist</div>
      </div>
      <div style="margin-bottom: 15px;">
        <label style="display: block; margin-bottom: 5px; font-weight: 500;">Filename (without extension):</label>
        <input type="text" id="__file_name__" value="extracted_data" 
               style="width: 100%; padding: 8px; border: 1px solid #ccc; border-radius: 4px; font-size: 14px;">
      </div>
      <div style="margin-bottom: 15px;">
        <div style="font-weight: 500; margin-bottom: 8px;">Mode:</div>
        <label style="display: block; margin: 5px 0; cursor: pointer;">
          <input type="radio" name="__file_mode__" value="append" checked style="margin-right: 8px;">
          <span>Append to existing file</span>
        </label>
        <label style="display: block; margin: 5px 0; cursor: pointer;">
          <input type="radio" name="__file_mode__" value="new" style="margin-right: 8px;">
          <span>Create new file (overwrite)</span>
        </label>
      </div>
      <div style="display: flex; justify-content: flex-end; gap: 8px;">
        <button id="__file_cancel__" style="padding: 8px 16px; background: #f44336; color: white; border: none; border-radius: 4px; cursor: pointer;">Cancel</button>
        <button id="__file_save__" style="padding: 8px 16px; background: #4CAF50; color: white; border: none; border-radius: 4px; cursor: pointer;">Save</button>
      </div>
    `;
    
    overlay.appendChild(dialog);
    document.body.appendChild(overlay);
    overlay.setAttribute('data-webai-dialog', 'true');
    
    const folderInput = document.getElementById('__file_folder__');
    const nameInput = document.getElementById('__file_name__');
    const cancelBtn = document.getElementById('__file_cancel__');
    const saveBtn = document.getElementById('__file_save__');
    
    const cleanup = () => overlay.remove();
    
    cancelBtn.onclick = (e) => {
      e.stopPropagation();
      cleanup();
      callback(null);
    };
    
    saveBtn.onclick = (e) => {
      e.stopPropagation();
      
      const folder = folderInput.value.trim();
      const filename = nameInput.value.trim();
      const mode = document.querySelector('input[name="__file_mode__"]:checked').value;
      
      if (!folder || !filename) {
        alert('Please enter both folder and filename');
        return;
      }
      
      cleanup();
      callback({
        formats: formats,
        folder: folder,
        filename: filename,
        mode: mode
      });
    };
  }

  function highlightExtracted(element, message) {
    const originalOutline = element.style.outline;
    const originalBg = element.style.backgroundColor;
    
    element.style.outline = '3px solid #4CAF50';
    element.style.backgroundColor = '#E8F5E9';
    
    if (typeof window.showHint === "function") {
      window.showHint(message);
    }
    
    setTimeout(() => {
      element.style.outline = originalOutline;
      element.style.backgroundColor = originalBg;
    }, 2000);
  }

  function removeExtractionMenu() {
    if (__CONTEXT_MENU__) {
      __CONTEXT_MENU__.remove();
      __CONTEXT_MENU__ = null;
    }
  }

  window.__cancelExtract = removeExtractionMenu;

  document.addEventListener('click', (e) => {
    if (__CONTEXT_MENU__ && !__CONTEXT_MENU__.contains(e.target)) {
      removeExtractionMenu();
    }
  }, true);
})();
"""


class DataExtractionPlugin:
    """
    Plugin that listens to data extraction events ('extract' and 'extract_table')
    and handles immediate background file persistence.
    """

    def attach(self, recorder: WebRecorder) -> None:
        """Attach plugin to WebRecorder instance by subscribing to event bus."""
        recorder.subscribe("extract", self.on_extract_event)
        recorder.subscribe("extract_table", self.on_extract_table_event)

    def on_extract_event(self, event_type: str, step: Step, payload: Dict[str, Any]) -> None:
        """Handle standard text/attribute extraction event."""
        save_options = payload.get("save_options")
        if save_options:
            try:
                self._save_extraction_immediately(step, save_options)
            except Exception as e:
                print(f" [DataExtractionPlugin] Error saving extraction immediately: {e}")

    def on_extract_table_event(self, event_type: str, step: Step, payload: Dict[str, Any]) -> None:
        """Handle table extraction event."""
        save_options = payload.get("save_options")
        if save_options:
            try:
                self._save_extraction_immediately(step, save_options)
            except Exception as e:
                print(f" [DataExtractionPlugin] Error saving table extraction immediately: {e}")

    def _save_extraction_immediately(self, step: Step, options: Dict[str, Any]) -> None:
        """Save single extraction to files immediately."""
        formats = options.get("formats", {})
        folder = options.get("folder", ".")
        filename = options.get("filename", "extracted_data")
        mode = options.get("mode", "append")

        if not os.path.exists(folder):
            try:
                os.makedirs(folder, exist_ok=True)
                print(f" [DataExtractionPlugin] Created folder: {folder}")
            except Exception as e:
                print(f" [DataExtractionPlugin] Cannot create folder ({e}), using current directory")
                folder = "."

        saved_files = []

        try:
            if formats.get("excel"):
                filepath = os.path.join(folder, f"{filename}.xlsx")
                self._save_to_excel_immediate(step, filepath, mode)
                saved_files.append(os.path.basename(filepath))
        except Exception as e:
            print(f" [DataExtractionPlugin] Excel save failed: {e}")

        try:
            if formats.get("word"):
                filepath = os.path.join(folder, f"{filename}.docx")
                self._save_to_word_immediate(step, filepath, mode)
                saved_files.append(os.path.basename(filepath))
        except Exception as e:
            print(f" [DataExtractionPlugin] Word save failed: {e}")

        try:
            if formats.get("txt"):
                filepath = os.path.join(folder, f"{filename}.txt")
                self._save_to_txt_immediate(step, filepath, mode)
                saved_files.append(os.path.basename(filepath))
        except Exception as e:
            print(f" [DataExtractionPlugin] Text save failed: {e}")

        if saved_files:
            print(f" [DataExtractionPlugin] Saved to: {', '.join(saved_files)}")

    def _save_to_excel_immediate(self, step: Step, filepath: str, mode: str) -> None:
        """Save extraction to Excel file."""
        try:
            import openpyxl
            from openpyxl.styles import Font
        except ImportError:
            print(" openpyxl not installed. Run: pip install openpyxl")
            return

        if mode == "new" or not os.path.exists(filepath):
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Extracted Data"

            headers = ["Variable", "Value", "Type", "Attribute", "URL", "Timestamp"]
            ws.append(headers)

            for cell in ws[1]:
                cell.font = Font(bold=True)
        else:
            wb = openpyxl.load_workbook(filepath)
            ws = wb.active

        ws.append([
            step.name,
            step.value or "",
            step.extract_type or "text",
            step.attribute_name or "",
            step.url or "",
            datetime.fromtimestamp(step.ts).strftime("%Y-%m-%d %H:%M:%S")
        ])

        for column in ws.columns:
            max_length = 0
            column_cells = [cell for cell in column]
            for cell in column_cells:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(cell.value)
                except Exception:
                    pass
            adjusted_width = min(max_length + 2, 50)
            ws.column_dimensions[column_cells[0].column_letter].width = adjusted_width

        wb.save(filepath)

    def _save_to_word_immediate(self, step: Step, filepath: str, mode: str) -> None:
        """Save extraction to Word file."""
        try:
            from docx import Document
        except ImportError:
            print(" python-docx not installed. Run: pip install python-docx")
            return

        if mode == "new" or not os.path.exists(filepath):
            doc = Document()
            doc.add_heading("Extracted Data", 0)
        else:
            doc = Document(filepath)

        doc.add_paragraph("━" * 60)
        doc.add_heading(f"Variable: {step.name}", level=2)
        doc.add_paragraph(f"Value: {step.value or '(empty)'}")
        doc.add_paragraph(f"Type: {step.extract_type or 'text'}")

        if step.attribute_name:
            doc.add_paragraph(f"Attribute: {step.attribute_name}")

        doc.add_paragraph(f"URL: {step.url or ''}")
        doc.add_paragraph(f"Timestamp: {datetime.fromtimestamp(step.ts).strftime('%Y-%m-%d %H:%M:%S')}")

        doc.save(filepath)

    def _save_to_txt_immediate(self, step: Step, filepath: str, mode: str) -> None:
        """Save extraction to text file."""
        write_mode = "w" if mode == "new" and os.path.exists(filepath) else "a"

        with open(filepath, write_mode, encoding="utf-8") as f:
            if write_mode == "w" or not os.path.exists(filepath) or os.path.getsize(filepath) == 0:
                f.write("=== Extracted Data Log ===\n")
                f.write(f"Started: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")

            f.write(f"{'─' * 60}\n")
            f.write(f"Variable: {step.name}\n")
            f.write(f"Value: {step.value or '(empty)'}\n")
            f.write(f"Type: {step.extract_type or 'text'}\n")

            if step.attribute_name:
                f.write(f"Attribute: {step.attribute_name}\n")

            f.write(f"URL: {step.url or ''}\n")
            f.write(f"Time: {datetime.fromtimestamp(step.ts).strftime('%Y-%m-%d %H:%M:%S')}\n\n")
