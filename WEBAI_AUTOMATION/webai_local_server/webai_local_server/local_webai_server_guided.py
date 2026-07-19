import asyncio
import hashlib
import json
import os
import re
import urllib.parse
import urllib.request
from pathlib import Path
from typing import Any, Dict, List, Optional
from datetime import datetime

import websockets


# ============================================================================
# PHASE 8.2: SAVE EXTRACTED DATA - Helpers
# ============================================================================

def _save_txt_extraction(filepath: str, mode: str, data: Dict[str, Any]) -> None:
    """
    Save the extracted data to a simple text file.
    
    This fulfills the extraction saving requirement where scraped data 
    (e.g., prices, text) is written persistently for the user.
    """
    import os
    from datetime import datetime
    
    write_mode = 'w' if mode == 'new' and os.path.exists(filepath) else 'a'
    
    with open(filepath, write_mode, encoding='utf-8') as f:
        if write_mode == 'w' or not os.path.exists(filepath) or os.path.getsize(filepath) == 0:
            f.write(f"=== Extracted Data Log ===\n")
            f.write(f"Started: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
        
        f.write(f"{'─' * 60}\n")
        f.write(f"Variable: {data.get('name', '')}\n")
        f.write(f"Value: {data.get('value', '(empty)')}\n")
        f.write(f"Type: {data.get('extract_type', 'text')}\n")
        
        if data.get('attribute_name'):
            f.write(f"Attribute: {data['attribute_name']}\n")
        
        f.write(f"URL: {data.get('url', '')}\n")
        f.write(f"Time: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")


def _save_excel_extraction(filepath: str, mode: str, data: Dict[str, Any]) -> None:
    """
    Save the extracted data to an Excel file using openpyxl.
    
    Creates a new workbook or appends to an existing one, formatting the 
    data into a structured table (Variable, Value, Type, Attribute, URL, Timestamp).
    """
    try:
        import openpyxl
        from openpyxl.styles import Font
    except ImportError:
        print("⚠️ openpyxl not installed. Run: pip install openpyxl")
        return
    
    import os
    from datetime import datetime
    
    if mode == 'new' or not os.path.exists(filepath):
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Extracted Data"
        
        headers = ['Variable', 'Value', 'Type', 'Attribute', 'URL', 'Timestamp']
        ws.append(headers)
        
        for cell in ws[1]:
            cell.font = Font(bold=True)
    else:
        wb = openpyxl.load_workbook(filepath)
        ws = wb.active
    
    ws.append([
        data.get('name', ''),
        data.get('value', ''),
        data.get('extract_type', 'text'),
        data.get('attribute_name', ''),
        data.get('url', ''),
        datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    ])
    
    for column in ws.columns:
        max_length = 0
        column = [cell for cell in column]
        for cell in column:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(cell.value)
            except:
                pass
        adjusted_width = min(max_length + 2, 50)
        ws.column_dimensions[column[0].column_letter].width = adjusted_width
    
    wb.save(filepath)


def _save_word_extraction(filepath: str, mode: str, data: Dict[str, Any]) -> None:
    """
    Save the extracted data to a Microsoft Word document using python-docx.
    
    Creates a new document or appends to an existing one, organizing the 
    data sequentially with headings and paragraphs.
    """
    try:
        from docx import Document
    except ImportError:
        print("⚠️ python-docx not installed. Run: pip install python-docx")
        return
    
    import os
    from datetime import datetime
    
    if mode == 'new' or not os.path.exists(filepath):
        doc = Document()
        doc.add_heading('Extracted Data', 0)
    else:
        doc = Document(filepath)
    
    doc.add_paragraph('━' * 60)
    doc.add_heading(f"Variable: {data.get('name', '')}", level=2)
    doc.add_paragraph(f"Value: {data.get('value', '(empty)')}")
    doc.add_paragraph(f"Type: {data.get('extract_type', 'text')}")
    
    if data.get('attribute_name'):
        doc.add_paragraph(f"Attribute: {data['attribute_name']}")
    
    doc.add_paragraph(f"URL: {data.get('url', '')}")
    doc.add_paragraph(f"Timestamp: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    
    doc.save(filepath)



def _task_hash(task_text: str) -> str:
    """
    Generate a short SHA-256 hash for a given natural language task.
    
    This hash is used as a cache key so that previously resolved task plans 
    can be reused without calling the LLM again.
    """
    task_text = (task_text or "").strip().encode("utf-8", errors="replace")
    return hashlib.sha256(task_text).hexdigest()[:12]


def _tid(task_id: str) -> str:
    """
    Shorten a task ID for logging purposes (first 8 characters).
    """
    return (task_id or "unknown")[:8]


def _fmt_target(target: dict) -> str:
    """
    Format a target dictionary (locator strategy) into a readable string.
    
    Converts a JSON-based locator object (e.g., `{"by": "label", "label": "Email"}`)
    into a human-readable format for logs and the LLM context.
    """
    if not isinstance(target, dict):
        return str(target)
    by = (target.get("by") or "").strip().lower()
    if by == "label":
        return f"label='{target.get('label','')}'"
    if by == "role":
        return f"role='{target.get('role','')}' name='{target.get('name','')}'"
    if by == "placeholder":
        return f"placeholder='{target.get('placeholder','')}'"
    if by == "text":
        return f"text='{target.get('text','')}'"
    return str(target)


class TaskDone(Exception):
    """
    Exception raised internally when the LLM outputs `action=done`.
    
    This signals the control loop that the current natural language task 
    has been fully completed and verified by the LLM.
    """
    pass

BASE_PROMPT = (
    "You are a careful browser automation planner.\n"
    "You MUST follow the allowed action schema exactly.\n"
    "Use the provided page context only (visible interactive elements list).\n"
    "Target preference order (MUST FOLLOW):\n"
    "  1) by='label' (most stable)\n"
    "  2) by='role' with role+name\n"
    "  3) by='placeholder'\n"
    "  4) by='text'\n"
    "If an element isn't visible, use scroll_page down and try again.\n"
    "Do NOT guess selectors. Do NOT invent elements.\n"
    "Do NOT output done until you have taken actions and verified success.\n"
    "Never use wait_text for URL verification. Use verify_url for URL checks.\n"
)

SUBGOALS = ("navigate", "act", "verify")


def build_subgoal_prompt(subgoal: str, base_prompt: str) -> str:
    """
    Construct a prompt tailored to a specific phase (subgoal) of task execution.
    
    The AI Brain executes tasks in three distinct phases to reduce hallucination:
    1. 'navigate' -> Focuses strictly on `goto` and `scroll_page`.
    2. 'act' -> Focuses strictly on interacting with the page (`click`, `type`, etc.).
    3. 'verify' -> Focuses strictly on verifying state (`wait_text`, `verify_url`).
    """
    if subgoal == "navigate":
        return (
            base_prompt
            + "\nSUBGOAL: NAVIGATION\n"
            + "Only plan actions to reach the correct page.\n"
            + "Allowed actions: goto, scroll_page.\n"
            + "Do NOT click or type unless required for navigation.\n"
            + "Use EXACT fields: goto uses 'url'. scroll_page uses 'target'. Do NOT use 'target' for goto and do NOT use 'direction' for scroll_page.\n"
        )

    if subgoal == "act":
        return (
            base_prompt
            + "\nSUBGOAL: ACTION\n"
            + "Only plan user interactions (click, type, select).\n"
            + "Do NOT verify success here.\n"
            + "If element not visible, scroll_page.\n"
            + "For scroll_page you MUST set target to one of: up/down/top/bottom. Never output null."
        )

    if subgoal == "verify":
        return (
            base_prompt
            + "\nSUBGOAL: VERIFICATION\n"
            + "Only verify success criteria.\n"
            + "Use wait_text for visible content.\n"
            + "Use verify_url for URL checks.\n"
            + "Do NOT navigate or click.\n"
        )

    raise ValueError(f"Unknown subgoal: {subgoal}")


def jdump(obj: Any) -> str:
    return json.dumps(obj, ensure_ascii=False)


def get_request_path(ws: Any) -> str:
    # websockets API differs by version
    if hasattr(ws, "path"):
        return ws.path
    req = getattr(ws, "request", None)
    if req and hasattr(req, "path"):
        return req.path
    return "/api"


def get_query_param(path: str, key: str) -> Optional[str]:
    parsed = urllib.parse.urlparse(path)
    qs = urllib.parse.parse_qs(parsed.query)
    vals = qs.get(key)
    return vals[0] if vals else None


def _env(name: str, default: str = "") -> str:
    val = os.getenv(name)
    return default if val is None or val == "" else val


async def ollama_chat(system: str, user: str) -> str:
    """
    Calls local Ollama chat API.
    Default endpoint: http://localhost:11434/api/chat
    Returns assistant message content.
    """
    url = _env("OLLAMA_URL", "http://localhost:11434/api/chat")
    model = _env("OLLAMA_MODEL", "llama3.1")

    payload = {
        "model": model,
        "stream": False,
        "messages": [
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ],
        "options": {
            "temperature": float(_env("OLLAMA_TEMPERATURE", "0.2")),
        },
    }

    data = jdump(payload).encode("utf-8")

    def _do_request() -> str:
        req = urllib.request.Request(url, data=data, headers={"Content-Type": "application/json"})
        with urllib.request.urlopen(req, timeout=60) as resp:
            raw = resp.read().decode("utf-8", errors="replace")
        obj = json.loads(raw)
        return obj.get("message", {}).get("content", "")

    return await asyncio.to_thread(_do_request)


ALLOWED_ACTIONS = """
Return ONLY a JSON array of actions. No markdown, no explanations.

CRITICAL: The 'by' field in targets MUST be one of: 'label', 'role', 'text', 'placeholder'
❌ NEVER use by='name' - it is NOT supported!

For visible text (like button labels, link text), use by='text'
For ARIA accessible names with roles, use by='role' with the 'name' parameter

Allowed actions (use targets; do NOT use CSS selectors):
1) {"action":"goto","url":"https://..."}
2) {"action":"click","target":{"by":"label","label":"Sign in","exact":false}}
3) {"action":"click","target":{"by":"role","role":"button","name":"Search","exact":false}}
4) {"action":"click","target":{"by":"text","text":"Sign in","exact":false}}
5) {"action":"type","target":{"by":"label","label":"Email","exact":false},"text":"..."}
6) {"action":"type","target":{"by":"placeholder","placeholder":"Search","exact":false},"text":"..."}
7) {"action":"type","target":{"by":"role","role":"textbox","name":"Search","exact":false},"text":"..."}
8) {"action":"select","target":{"by":"label","label":"Country","exact":false},"option_text":"India"}
9) {"action":"select_search","target":{"by":"label","label":"City","exact":false},"query":"che","option_text":"Chennai"}
10) {"action":"scroll_page","target":"down"}   // up/down/top/bottom
11) {"action":"press_key","key":"Enter"}       // any Playwright key, e.g. Enter, Tab, Control+A
12) {"action":"wait_text","text":"Welcome","timeout_ms":10000}
13) {"action":"wait_role","role":"button","name":"Submit","exact":false,"timeout_ms":10000}
14) {"action":"done","summary":"..."}
15) {"action":"verify_url","contains":"contact-us"}
16) {"action":"verify_element","target":{"by":"text","text":"Get in touch","exact":false},"timeout_ms":10000}

""".strip()


def _compact_context(url: str, title: str, elements: List[Dict[str, Any]]) -> str:
    """Compact “UI inventory” to reduce hallucinations."""
    lines = [f"URL: {url}", f"TITLE: {title}", "VISIBLE INTERACTIVE ELEMENTS (partial):"]
    for i, el in enumerate((elements or [])[:60]):
        desc = {
            "tag": el.get("tag", ""),
            "role": el.get("role", ""),
            "type": el.get("type", ""),
            "name": el.get("name", ""),
            "placeholder": el.get("placeholder", ""),
            "ariaLabel": el.get("ariaLabel", ""),
            "text": el.get("text", ""),
        }
        lines.append(f"{i+1}. {jdump(desc)}")
    return "\n".join(lines)


def _extract_json_array(text: str) -> List[Dict[str, Any]]:
    """Strict parser: expects a JSON array; extracts first [...] if extra text exists."""
    text = (text or "").strip()

    try:
        obj = json.loads(text)
        if isinstance(obj, list):
            return obj
    except Exception:
        pass

    s = text.find("[")
    e = text.rfind("]")
    if s != -1 and e != -1 and e > s:
        obj = json.loads(text[s : e + 1])
        if isinstance(obj, list):
            return obj

    raise RuntimeError(f"LLM returned non-JSON plan: {text[:300]}")


# -----------------------------
# Selector string -> target normalization (prevents Unsupported type target: {})
# -----------------------------
_SELECTOR_RE = re.compile(
    r"by='(?P<by>\w+)'(?:\s+with\s+(?P<extra>[^:]+))?:\s*'(?P<value>[^']*)'",
    re.IGNORECASE,
)

def selector_to_target(selector: str) -> Dict[str, Any]:
    """
    Parse legacy selector strings like:
      by='text': 'OK (Enter)'
      by='role' with role+name: 'Search'
    into a structured target dict used by click/type actions.
    """
    s = (selector or "").strip()
    m = _SELECTOR_RE.search(s)
    if not m:
        return {}

    by = (m.group("by") or "").strip().lower()
    extra = (m.group("extra") or "").strip().lower()
    value = (m.group("value") or "").strip()

    if by == "text":
        return {"by": "text", "text": value, "exact": False}
    if by == "label":
        return {"by": "label", "label": value, "exact": False}
    if by == "placeholder":
        return {"by": "placeholder", "placeholder": value, "exact": False}
    if by == "role":
        # If role isn't specified, default to textbox (safe for typing)
        role = "textbox"
        if "button" in extra:
            role = "button"
        return {"by": "role", "role": role, "name": value, "exact": False}

    return {}


# -----------------------------
# Task normalization
# -----------------------------
URL_RE = re.compile(r"https?://[^\s)>\"']+")


def _is_already_structured(task: str) -> bool:
    t = (task or "").lower()
    return ("goal:" in t) or ("requirements:" in t) or ("success criteria:" in t)


def _extract_urls(task: str) -> List[str]:
    return URL_RE.findall(task or "")


def normalize_task(task: str) -> str:
    """Convert a short task into Goal/Requirements/Success Criteria form."""
    task = (task or "").strip()
    if not task or _is_already_structured(task):
        return task

    urls = _extract_urls(task)
    primary_url = urls[0] if urls else ""

    lower = task.lower()
    wants_search = any(k in lower for k in ["search", "google", "type", "enter", "query"])
    wants_click = any(k in lower for k in ["click", "tap"])
    wants_nav = bool(primary_url) or any(k in lower for k in ["go to", "open", "launch"])

    lines: List[str] = []
    if primary_url:
        lines.append(f"Open: {primary_url}")
        if len(urls) > 1:
            lines.append(f"Other URLs mentioned: {', '.join(urls[1:])}")
        lines.append("")

    lines.append("Goal:")
    if wants_search:
        lines.append("- Complete the requested search/action described below.")
    elif wants_click:
        lines.append("- Navigate and perform the requested click/navigation action.")
    elif wants_nav:
        lines.append("- Open the site and complete the requested action.")
    else:
        lines.append("- Complete the task described below.")

    lines.append("")
    lines.append("Requirements:")
    lines.append("- Wait until the page is fully loaded before interacting.")
    lines.append("- Prefer stable targeting: label first, then role+name, then placeholder, then visible text.")
    lines.append("- If an element is not visible, scroll down and try again.")
    lines.append("- Avoid guessing; use visible interactive elements in context.")
    lines.append("")
    lines.append("Task details (user request):")
    lines.append(f"- {task}")
    lines.append("")
    lines.append("Success criteria:")
    lines.append("- The requested action is completed and observable (URL/title/content changes appropriately).")
    lines.append("- If navigation is expected, confirm using URL change and/or visible page content.")
    lines.append("")
    lines.append("Finish with done only after all success criteria are met.")
    return "\n".join(lines)


# -----------------------------
# Success rules extraction
# -----------------------------
QUOTED_RE = re.compile(r"""["']([^"']{2,})["']""")
URL_CONTAINS_RE = re.compile(r"""url\s+contains\s+["']([^"']+)["']""", re.IGNORECASE)


def _infer_task_type(task: str) -> str:
    t = (task or "").lower()
    if any(k in t for k in ["select", "dropdown", "autocomplete", "choose", "pick an option"]):
        return "dropdown"
    if any(k in t for k in ["fill", "form", "submit", "enter details", "sign up", "register"]):
        return "form"
    if any(k in t for k in ["search", "query", "type", "press enter", "google"]):
        return "search"
    if any(k in t for k in ["navigate", "go to", "open", "launch", "click", "contact us", "login"]):
        return "navigation"
    return "generic"

import re

_VERIFY_PAGE_TEXT_RE = re.compile(r'^\s*Verify\s+the\s+page\s+contains\s+text\s+"(.+?)"\s*$', re.IGNORECASE)
_VERIFY_ELEMENT_RE = re.compile(r'^\s*Verify\s+element\s+"(.+?)"\s+is\s+visible\s*$', re.IGNORECASE)
_VERIFY_URL_CONTAINS_RE = re.compile(r'^\s*Verify\s+URL\s+contains\s+"(.+?)"\s*$', re.IGNORECASE)

def extract_success_expectations(task: str) -> Dict[str, Any]:
    """
    IMPORTANT CHANGE:
    - Only treat explicit VERIFY lines as success expectations.
    - Do NOT treat all quoted strings (like input values) as expected texts.
    """
    task = task or ""
    expected_texts: List[str] = []
    expected_url_substrings: List[str] = []

    for line in task.splitlines():
        line = line.strip()
        if not line:
            continue

        m = _VERIFY_PAGE_TEXT_RE.match(line)
        if m:
            expected_texts.append(m.group(1).strip())
            continue

        m = _VERIFY_ELEMENT_RE.match(line)
        if m:
            expected_texts.append(m.group(1).strip())
            continue

        m = _VERIFY_URL_CONTAINS_RE.match(line)
        if m:
            expected_url_substrings.append(m.group(1).strip())
            continue

    # Dedup (case-insensitive)
    def _dedup(seq):
        seen = set()
        out = []
        for s in seq:
            k = s.lower()
            if k not in seen:
                seen.add(k)
                out.append(s)
        return out

    return {
        "task_type": _infer_task_type(task),
        "expected_url_substrings": _dedup([s for s in expected_url_substrings if s]),
        "expected_texts": _dedup([s for s in expected_texts if s]),
    }


def url_meets_expectations(url: str, expected_substrings: List[str]) -> bool:
    if not expected_substrings:
        return False
    u = (url or "").lower()
    return any(sub.lower() in u for sub in expected_substrings)


def plan_requires_strict_success(task: str) -> bool:
    t = (task or "").lower()
    return ("success criteria:" in t) or ("url contains" in t) or ("look for the words" in t)


def _shorten(s: str, n: int = 120) -> str:
    s = (s or "").strip()
    return s if len(s) <= n else s[:n] + "..."


# -----------------------------
# Retry helpers
# -----------------------------
def _looks_like_not_found(err: str) -> bool:
    e = (err or "").lower()
    return any(k in e for k in [
        "not found",
        "no node found",
        "could not find",
        "timeout",
        "strict mode violation",
        "element is not",
        "has been detached",
        "intercepted",
        "another element",
    ])


def _looks_like_navigation_issue(err: str) -> bool:
    e = (err or "").lower()
    return any(k in e for k in ["navigation", "net::", "frame was detached", "target closed"])


def _format_last_errors(errors: List[str], limit: int = 3) -> str:
    if not errors:
        return ""
    recent = errors[-limit:]
    lines = ["LAST ERRORS (most recent last):"]
    for i, e in enumerate(recent, 1):
        lines.append(f"{i}. {e}")
    return "\n".join(lines)


# -----------------------------
# Confidence + verification
# -----------------------------
def compute_confidence(progress_ok: bool, strict: bool, url_ok: bool, text_ok: bool, failures: int) -> int:
    score = 40
    if progress_ok:
        score += 25
    if strict:
        if url_ok:
            score += 20
        if text_ok:
            score += 20
        if not url_ok and not text_ok:
            score -= 25
    else:
        score += 10
    score -= min(30, failures * 8)
    return max(0, min(100, score))


async def verify_texts(send_command, expected_texts: List[str]) -> bool:
    if not expected_texts:
        return False
    for txt in expected_texts[:3]:
        try:
            await send_command("waitForText", {"text": txt, "timeoutMs": 2500})
            return True
        except Exception:
            continue
    return False


# -----------------------------
# Plan cache (record/replay)
# -----------------------------
CACHE_PATH = Path(_env("PLAN_CACHE_FILE", "plan_cache.json"))


def _cache_key(url: str, task_text: str) -> str:
    url = (url or "").strip()
    task_text = (task_text or "").strip().lower()
    raw = f"{url}||{task_text}"  # include full task, not only first line
    return hashlib.sha256(raw.encode("utf-8")).hexdigest()[:16]



def _load_cache() -> Dict[str, Any]:
    if not CACHE_PATH.exists():
        return {}
    try:
        return json.loads(CACHE_PATH.read_text(encoding="utf-8"))
    except Exception:
        return {}


def _save_cache(cache: Dict[str, Any]) -> None:
    CACHE_PATH.write_text(json.dumps(cache, ensure_ascii=False, indent=2), encoding="utf-8")


def cache_get_plan(url: str, task_text: str) -> Optional[List[Dict[str, Any]]]:
    cache = _load_cache()
    key = _cache_key(url, task_text)
    entry = cache.get(key)
    if not entry:
        return None
    plan = entry.get("plan")
    return plan if isinstance(plan, list) else None


def cache_put_plan(url: str, task_text: str, plan: List[Dict[str, Any]]) -> None:
    cache = _load_cache()
    key = _cache_key(url, task_text)
    cache[key] = {
        "url": url,
        "task": task_text.splitlines()[0][:200],
        "plan": plan,
        "updated_at": int(asyncio.get_event_loop().time()),
    }
    _save_cache(cache)


# -----------------------------
# Templates / Skill registry
# -----------------------------
def build_system_prompt(task_text: str, expect: Dict[str, Any]) -> str:
    task_type = (expect.get("task_type") or "generic").strip().lower()

    NAV_TEMPLATE = (
        "NAVIGATION TEMPLATE:\n"
        "- Click the relevant link/button and confirm navigation.\n"
        "- After clicking, verify navigation by URL/title/content change.\n"
        "- For hyperlinks/menu items, prefer role='link' with name, then by='text'.\n"
    )

    SEARCH_TEMPLATE = (
        "SEARCH TEMPLATE:\n"
        "- Click textbox -> type query -> press_key Enter.\n"
        "- Verify success by URL/content change.\n"
    )

    FORM_TEMPLATE = (
        "FORM TEMPLATE:\n"
        "- Fill fields then submit.\n"
        "- Prefer label targeting for inputs.\n"
        "- Verify success by confirmation text or navigation.\n"
    )

    DROPDOWN_TEMPLATE = (
        "DROPDOWN TEMPLATE:\n"
        "- Use selectSmart or selectSearchSmart.\n"
        "- Verify selected value using wait_text.\n"
    )

    template = NAV_TEMPLATE
    if task_type == "search":
        template = SEARCH_TEMPLATE
    elif task_type == "form":
        template = FORM_TEMPLATE
    elif task_type == "dropdown":
        template = DROPDOWN_TEMPLATE

    expected_urls = expect.get("expected_url_substrings", [])
    expected_texts = expect.get("expected_texts", [])

    hints = "EXPECTATION HINTS:\n"
    if expected_urls:
        hints += f"- Expected URL contains: {expected_urls}\n"
    if expected_texts:
        hints += f"- Expected text (use wait_text): {expected_texts[:3]}\n"

    return BASE_PROMPT + "\n" + template + "\n" + hints + "\n\n" + ALLOWED_ACTIONS

OPEN_LINE_RE = re.compile(r"^\s*Open:\s*(https?://\S+)\s*$", re.IGNORECASE | re.MULTILINE)

def extract_primary_open_url(task_text: str) -> Optional[str]:
    """
    Pulls the first Open: URL from the normalized task text.
    Example: "Open: https://sujaypublicschool.com/"
    """
    if not task_text:
        return None
    m = OPEN_LINE_RE.search(task_text)
    return m.group(1).strip() if m else None

# -----------------------------
# Observable progress heuristic
# -----------------------------
def _has_progress(prev_url: Optional[str], prev_title: Optional[str], url: str, title: str) -> bool:
    if not prev_url and not prev_title:
        return False
    if prev_url and url != prev_url:
        return True
    if prev_title and title != prev_title:
        return True
    return False


async def handle_client(ws: Any):
    """
    Core WebSocket connection handler for incoming automation clients.
    
    This function implements the "AI Brain" component of the 3-tier architecture. 
    It receives raw DOM snapshots from the Playwright browser robot, constructs 
    contextual prompts based on the current step, queries the local Ollama LLM, 
    and sends structured execution commands back to the browser.
    
    It manages the lifecycle of a task execution, including subgoal transitions 
    (navigate -> act -> verify), API server logging, and fallback recovery.
    """
    did_any_action = False

    path = get_request_path(ws)
    client_key = get_query_param(path, "key")
    print(f"✅ Client connected: {path} | key={'(none)' if not client_key else client_key}")

    task_id: Optional[str] = None
    command_index = 0

    actions_attempted = 0
    failures = 0
    last_errors: List[str] = []

    async def recv_json() -> Dict[str, Any]:
        raw = await ws.recv()
        if isinstance(raw, (bytes, bytearray)):
            raw = raw.decode("utf-8", errors="replace")
        return json.loads(raw)

    async def send_json(obj: Dict[str, Any]) -> None:
        await ws.send(json.dumps(obj))

    async def send_command(name: str, arguments: Dict[str, Any]) -> Any:
        nonlocal command_index, task_id

        await send_json({
            "type": "command-request",
            "taskId": task_id,
            "index": command_index,
            "name": name,
            "arguments": arguments or {},
        })
        command_index += 1

        while True:
            resp = await recv_json()
            if resp.get("type") == "command-response" and resp.get("taskId") == task_id:
                result = resp.get("result")
                if result and result != "null":
                    try:
                        return json.loads(result)
                    except Exception:
                        return result
                return None


    def normalize_action(act: Dict[str, Any]) -> Dict[str, Any]:
        """
        Make LLM output tolerant to common field-name variations.
        Converts:
        goto: target -> url
        scroll_page: direction -> target
        press_key: key variants
        Also converts:
        selector/element -> structured target (prevents Unsupported type target: {})
        """
        if not isinstance(act, dict):
            return act

        kind = (act.get("action") or "").strip().lower()

        # Legacy/alternate schema support:
        # Some plans return a string selector instead of a structured target.
        if not act.get("target"):
            if act.get("selector"):
                act["target"] = selector_to_target(act.get("selector", ""))
            elif act.get("element"):
                # best-effort: treat element as visible text
                act["target"] = {"by": "text", "text": str(act.get("element", "")), "exact": False}

        # SAFETY: Convert invalid by='name' to by='text'
        # This is a fallback if the LLM ignores our prompt instructions
        if kind in ("click", "type") and isinstance(act.get("target"), dict):
            t = act["target"]
            by_value = (t.get("by") or "").strip().lower()
            if by_value == "name":
                # Convert by='name' to by='text'
                print(f"[WARN] LLM used invalid by='name', auto-converting to by='text': {t}")
                act["target"] = {
                    "by": "text",
                    "text": t.get("name", ""),
                    "exact": bool(t.get("exact", False))
                }

        # If role target is missing role, default to textbox for typing
        if kind == "type" and isinstance(act.get("target"), dict):
            t = act["target"]
            if (t.get("by") or "").strip().lower() == "role" and not t.get("role"):
                t["role"] = "textbox"

        # goto: allow {target:"https://..."} in addition to {url:"https://..."}
        if kind == "goto":
            if not act.get("url") and act.get("target"):
                # allow old field usage
                if isinstance(act.get("target"), str):
                    act["url"] = act.get("target")
            # also allow accidental key name "href"
            if not act.get("url") and act.get("href"):
                act["url"] = act.get("href")

        # scroll_page: allow {direction:"down"} in addition to {target:"down"}
        if kind == "scroll_page":
            if not act.get("target") and act.get("direction"):
                act["target"] = act["direction"]
            if act.get("target") in (None, "", "null"):
                act["target"] = "down"

        # press_key: allow common variants
        if kind == "press_key":
            if not act.get("key") and act.get("button"):
                act["key"] = act.get("button")

        return act

    async def exec_action(act: Dict[str, Any]) -> None:
        nonlocal did_any_action

        kind = (act.get("action") or "").strip().lower()
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S.%f")[:-3]
        print(f"\n{log_prefix} [{timestamp}] ▶ ACTION: {kind.upper()}")

        if kind == "goto":
            url = (act.get("url") or "").strip()
            if not url:
                raise RuntimeError(f"goto missing url: {act}")
            print(f"{log_prefix} 🌐 Navigating to: {url}")
            await send_command("navigate", {"url": url})
            print(f"{log_prefix} ✅ Navigation complete")
            did_any_action = True
            return

        if kind == "scroll_page":
            direction = act.get("target", "down")
            print(f"{log_prefix} \u21d5\ufe0f  Scrolling page: {direction}")
            await send_command("scrollPage", {"target": direction})
            print(f"{log_prefix} \u2705 Scroll successful")
            did_any_action = True
            return

        if kind == "press_key":
            key = act.get("key", "Enter")
            print(f"{log_prefix} \u23ce  Pressing key: {key}")
            await send_command("pressKey", {"key": key})
            print(f"{log_prefix} \u2705 Key press successful")
            did_any_action = True
            return

        if kind == "wait_text":
            await send_command("waitForText", {"text": act.get("text", ""), "timeoutMs": int(act.get("timeout_ms", 10000))})
            return

        if kind == "wait_role":
            await send_command("waitForRole", {
                "role": act.get("role", ""),
                "name": act.get("name", ""),
                "exact": bool(act.get("exact", False)),
                "timeoutMs": int(act.get("timeout_ms", 10000)),
            })
            return

        if kind == "click":
            target = act.get("target") or {}
            by = (target.get("by") or "").strip().lower()

            if by == "label":
                label = target.get("label", "")
                print(f"{log_prefix} 🖱️  Clicking element by label: '{label}'")
                await send_command("clickByLabel", {"label": label, "exact": bool(target.get("exact", False))})
                print(f"{log_prefix} ✅ Click successful")
                did_any_action = True
                return
            if by == "role":
                role = target.get("role", "")
                name = target.get("name", "")
                print(f"{log_prefix} 🖱️  Clicking {role} with name: '{name}'")
                await send_command("clickByRole", {"role": role, "name": name, "exact": bool(target.get("exact", False))})
                print(f"{log_prefix} ✅ Click successful")
                did_any_action = True
                return
            if by == "text":
                text = target.get("text", "")
                print(f"{log_prefix} 🖱️  Clicking element with text: '{text}'")
                await send_command("clickByText", {"text": text, "exact": bool(target.get("exact", False))})
                print(f"{log_prefix} ✅ Click successful")
                did_any_action = True
                return
            
            # Helpful error message
            if by == "name":
                raise RuntimeError(
                    f"Invalid click target: by='name' is NOT supported. "
                    f"Use by='text' for visible text or by='role' with name parameter. "
                    f"Target: {target}"
                )
            
            raise RuntimeError(
                f"Unsupported click target (by='{by}'). "
                f"Supported values: 'label', 'role', 'text', 'placeholder'. "
                f"Target: {target}"
            )

        if kind == "type":
            target = act.get("target") or {}
            text = act.get("text", "")
            by = (target.get("by") or "").strip().lower()

            if by == "id":
                element_id = target.get("value", "")
                print(f"{log_prefix} ⌨️  Typing into element with id '{element_id}': '{text}'")
                await send_command("typeById", {"id": element_id, "text": text})
                print(f"{log_prefix} ✅ Type successful")
                did_any_action = True
                return
            if by == "name":
                name_attr = target.get("value", "")
                print(f"{log_prefix} ⌨️  Typing into element with name '{name_attr}': '{text}'")
                await send_command("typeByName", {"name": name_attr, "text": text})
                print(f"{log_prefix} ✅ Type successful")
                did_any_action = True
                return
            if by == "css":
                css = target.get("value", "")
                print(f"{log_prefix} ⌨️  Typing into element with CSS '{css[:50]}': '{text}'")
                await send_command("typeByCSS", {"css": css, "text": text})
                print(f"{log_prefix} ✅ Type successful")
                did_any_action = True
                return
            if by == "xpath":
                xpath = target.get("value", "")
                print(f"{log_prefix} ⌨️  Typing into element with XPath: '{text}'")
                await send_command("typeByXPath", {"xpath": xpath, "text": text})
                print(f"{log_prefix} ✅ Type successful")
                did_any_action = True
                return
            if by == "aria-label":
                aria_label = target.get("value", "")
                print(f"{log_prefix} ⌨️  Typing into element with aria-label '{aria_label}': '{text}'")
                await send_command("typeByAriaLabel", {"aria_label": aria_label, "text": text})
                print(f"{log_prefix} ✅ Type successful")
                did_any_action = True
                return
            if by == "label":
                label = target.get("value", "") or target.get("label", "")
                print(f"{log_prefix} ⌨️  Typing into field '{label}': '{text}'")
                await send_command("typeByLabel", {"label": label, "text": text, "exact": bool(target.get("exact", False))})
                print(f"{log_prefix} ✅ Type successful")
                did_any_action = True
                return
            if by == "placeholder":
                placeholder = target.get("value", "") or target.get("placeholder", "")
                print(f"{log_prefix} ⌨️  Typing into field with placeholder '{placeholder}': '{text}'")
                await send_command("typeByPlaceholder", {"placeholder": placeholder, "text": text, "exact": bool(target.get("exact", False))})
                print(f"{log_prefix} ✅ Type successful")
                did_any_action = True
                return
            if by == "role":
                role = target.get("role", "") or "textbox"
                name = target.get("name", "")
                print(f"{log_prefix} ⌨️  Typing into {role} '{name}': '{text}'")
                await send_command("typeByRole", {"role": role, "name": name, "text": text, "exact": bool(target.get("exact", False))})
                print(f"{log_prefix} ✅ Type successful")
                did_any_action = True
                return
            raise RuntimeError(f"Unsupported type target: {target}")

        if kind == "select":
            target = act.get("target") or {}
            option_text = act.get("option_text", "")
            await send_command("selectSmart", {"target": target, "optionText": option_text})
            did_any_action = True
            return

        if kind == "select_search":
            target = act.get("target") or {}
            query = act.get("query", "")
            option_text = act.get("option_text", "")
            await send_command("selectSearchSmart", {"target": target, "query": query, "optionText": option_text})
            did_any_action = True
            return

        if kind == "verify_url":
            await send_command("verifyUrl", {"contains": act.get("contains", "")})
            return

        if kind == "verify_element":
            # Option 1: Verify by TEXT only (waitForText)
            target = act.get("target") or {}
            timeout_ms = int(act.get("timeout_ms", 10000))
            by = (target.get("by") or "").strip().lower()
            if by == "text":
                await send_command("waitForText", {"text": target.get("text", ""), "timeoutMs": timeout_ms})
                return
            # fallback: role
            if by == "role":
                await send_command("waitForRole", {
                    "role": target.get("role", "") or "textbox",
                    "name": target.get("name", ""),
                    "exact": bool(target.get("exact", False)),
                    "timeoutMs": timeout_ms,
                })
                return
            raise RuntimeError(f"Unsupported verify_element target: {target}")

        # NEW: Fallback strategy actions
        if kind == "typewithfallback":
            locators = act.get("locators", [])
            text = act.get("text", "")
            print(f"{log_prefix} 🔄 Typing with fallback: '{text}' ({len(locators)} locators)")
            await send_command("typeWithFallback", {"locators": locators, "text": text})
            print(f"{log_prefix} ✅ Type successful")
            did_any_action = True
            return

        if kind == "clickwithfallback":
            locators = act.get("locators", [])
            print(f"{log_prefix} 🔄 Clicking with fallback ({len(locators)} locators)")
            await send_command("clickWithFallback", {"locators": locators})
            print(f"{log_prefix} ✅ Click successful")
            did_any_action = True
            return

        if kind == "extract":
            key = act.get("name", "extracted_value")  # FIXED: Use 'name' field (variable name in Step)
            locators = act.get("locators", [])
            extract_type = act.get("extract_type", "text")
            attribute_name = act.get("attribute_name")
            
            print(f"\n{log_prefix} 🔍 EXTRACTING data:")
            print(f"{log_prefix}    Variable: '{key}'")
            print(f"{log_prefix}    Type: {extract_type}")
            if attribute_name:
                print(f"{log_prefix}    Attribute: {attribute_name}")
            print(f"{log_prefix}    Trying {len(locators)} locators...")
            
            await send_command("extractData", {
                "key": key,
                "locators": locators,
                "extractType": extract_type,
                "attributeName": attribute_name
            })
            
            print(f"{log_prefix} ✅ Extraction complete: '{key}'")
            
            # NEW: Phase 8.2 - Handle save during playback if configured
            save_options = act.get("save_options")
            if save_options:
                formats = save_options.get('formats', {})
                folder = save_options.get('folder', '.')
                filename = save_options.get('filename', 'extracted_data')
                mode = save_options.get('mode', 'append')
                
                # Auto-create folder if needed
                import os
                if not os.path.exists(folder):
                    try:
                        os.makedirs(folder, exist_ok=True)
                        print(f"✅ Created folder: {folder}")
                    except Exception as e:
                        print(f"⚠️ Cannot create folder ({e}), using current directory")
                        folder = '.'
                
                # Prepare data for saving
                save_data = {
                    'name': key,
                    'value': act.get("value", ""),
                    'extract_type': extract_type,
                    'attribute_name': attribute_name,
                    'url': act.get("url", "")
                }
                
                saved_files = []
                
                try:
                    if formats.get('txt'):
                        filepath = os.path.join(folder, f"{filename}.txt")
                        _save_txt_extraction(filepath, mode, save_data)
                        saved_files.append(os.path.basename(filepath))
                except Exception as e:
                    print(f"⚠️ Text save failed: {e}")
                
                try:
                    if formats.get('excel'):
                        filepath = os.path.join(folder, f"{filename}.xlsx")
                        _save_excel_extraction(filepath, mode, save_data)
                        saved_files.append(os.path.basename(filepath))
                except Exception as e:
                    print(f"⚠️ Excel save failed: {e}")
                
                try:
                    if formats.get('word'):
                        filepath = os.path.join(folder, f"{filename}.docx")
                        _save_word_extraction(filepath, mode, save_data)
                        saved_files.append(os.path.basename(filepath))
                except Exception as e:
                    print(f"⚠️ Word save failed: {e}")
                
                if saved_files:
                    print(f"💾 Saved to: {', '.join(saved_files)}")
            
            did_any_action = True
            return

        if kind == "extract_table":
            # NEW: Phase 8.3 - Table extraction
            key = act.get("name", "table_data")
            table_config = act.get("table_config", {})
            locators = act.get("locators", [])
            save_options = act.get("save_options")
            
            print(f"\n{log_prefix} 📊 EXTRACTING TABLE:")
            print(f"{log_prefix}    Variable: '{key}'")
            print(f"{log_prefix}    Columns: {table_config.get('columns', [])}")
            
            pagination_config = table_config.get('pagination', {})
            if pagination_config.get('enabled'):
                print(f"{log_prefix}    Pagination: Enabled (max {pagination_config.get('max_pages', 10)} pages)")
                print(f"{log_prefix}    Wait: {pagination_config.get('wait_per_page', 2.0)}s, Timeout: {pagination_config.get('page_timeout', 10.0)}s, Retries: {pagination_config.get('retry_attempts', 3)}")
            
            # Extract table data using JavaScript evaluation via WebSocket
            table_selector = table_config.get('table_selector', '')
            columns = table_config.get('columns', [])
            column_indices = table_config.get('column_indices', [])
            max_pages = pagination_config.get('max_pages', 10) if pagination_config.get('enabled') else 1
            wait_per_page = pagination_config.get('wait_per_page', 2.0)
            page_timeout = pagination_config.get('page_timeout', 10.0)
            retry_attempts = pagination_config.get('retry_attempts', 3)
            
            # Send table extraction command to client
            result = await send_command("extractTableData", {
                "tableSelector": table_selector,
                "columns": columns,
                "columnIndices": column_indices,
                "maxPages": max_pages,
                "waitPerPage": wait_per_page,
                "pageTimeout": page_timeout,
                "retryAttempts": retry_attempts
            })
            
            print(f"{log_prefix} ✅ Extracted {result.get('row_count', 0)} rows")
            
            # Save table data if save options configured
            if save_options:
                formats = save_options.get('formats', {})
                folder = save_options.get('folder', '.')
                filename = save_options.get('filename', 'table_data')
                mode = save_options.get('mode', 'append')
                
                # Auto-create folder if needed
                import os
                if not os.path.exists(folder):
                    try:
                        os.makedirs(folder, exist_ok=True)
                        print(f"✅ Created folder: {folder}")
                    except Exception as e:
                        print(f"⚠️ Cannot create folder ({e}), using current directory")
                        folder = '.'
                
                # Get extracted table data from result
                table_data = result.get('data', [])
                
                if table_data:
                    saved_files = []
                    
                    try:
                        if formats.get('excel'):
                            import pandas as pd
                            filepath = os.path.join(folder, f"{filename}.xlsx")
                            df = pd.DataFrame(table_data)
                            
                            if mode == 'append' and os.path.exists(filepath):
                                # Append to existing file
                                existing_df = pd.read_excel(filepath)
                                df = pd.concat([existing_df, df], ignore_index=True)
                            
                            df.to_excel(filepath, index=False)
                            saved_files.append(os.path.basename(filepath))
                    except Exception as e:
                        print(f"⚠️ Excel save failed: {e}")
                    
                    try:
                        if formats.get('csv'):
                            import pandas as pd
                            filepath = os.path.join(folder, f"{filename}.csv")
                            df = pd.DataFrame(table_data)
                            
                            if mode == 'append' and os.path.exists(filepath):
                                df.to_csv(filepath, mode='a', header=False, index=False)
                            else:
                                df.to_csv(filepath, index=False)
                            
                            saved_files.append(os.path.basename(filepath))
                    except Exception as e:
                        print(f"⚠️ CSV save failed: {e}")
                    
                    try:
                        if formats.get('txt'):
                            filepath = os.path.join(folder, f"{filename}.txt")
                            with open(filepath, 'a' if mode == 'append' else 'w') as f:
                                for row in table_data:
                                    f.write(str(row) + '\n')
                            saved_files.append(os.path.basename(filepath))
                    except Exception as e:
                        print(f"⚠️ Text save failed: {e}")
                    
                    if saved_files:
                        print(f"💾 Saved table to: {', '.join(saved_files)}")
            
            did_any_action = True
            return

        if kind == "wait":
            # NEW: Handle explicit wait/delay (Phase 9.1)
            seconds = act.get("value")
            reason = act.get("name")
            
            try:
                wait_time = float(seconds)
                if wait_time < 1 or wait_time > 60:
                    print(f"{log_prefix} ⚠️ Invalid wait time: {wait_time}s (skipping)")
                    return
                
                start_time = datetime.now()
                timestamp = start_time.strftime("%Y-%m-%d %H:%M:%S")
                
                if reason:
                    print(f"\n[{timestamp}] {log_prefix} ⏱️ Waiting {wait_time} seconds ({reason})...")
                else:
                    print(f"\n[{timestamp}] {log_prefix} ⏱️ Waiting {wait_time} seconds...")
                
                await asyncio.sleep(wait_time)
                
                end_time = datetime.now()
                end_timestamp = end_time.strftime("%Y-%m-%d %H:%M:%S")
                elapsed = (end_time - start_time).total_seconds()
                
                print(f"[{end_timestamp}] {log_prefix} ✅ Wait complete (actual: {elapsed:.1f}s)")
                did_any_action = True
                return
            except (ValueError, TypeError) as e:
                print(f"{log_prefix} ⚠️ Invalid wait value: {seconds}")
                return

        if kind == "done":
            # Only accept done if we actually executed something.
            if not did_any_action:
                # Ignore premature done and let the loop continue.
                return
            await send_json({"type": "task-complete", "taskId": task_id, "success": True})
            raise TaskDone()

        raise RuntimeError(f"Unsupported action: {act}")

    try:
        start_msg = await recv_json()
        if start_msg.get("type") != "task-start":
            await send_json({
                "type": "task-complete",
                "taskId": start_msg.get("taskId", "unknown"),
                "success": False,
                "error": "Expected task-start first",
            })
            return

        task_id = start_msg["taskId"]
        raw_task_text = start_msg.get("task", "")
        if _env("TASK_NORMALIZATION", "1") == "1":
            task_text = normalize_task(raw_task_text)
            print("🧩 Normalized task:\n" + task_text)
        else:
            task_text = raw_task_text
        options = start_msg.get("options") or {}
        task_hash = _task_hash(task_text)
        log_prefix = f"[task={_tid(task_id)} hash={task_hash}]"
        print(f"{log_prefix} ✅ task-start received")
        if options.get("recorded_steps"):
            try:
                rs_len = len(options.get("recorded_steps") or [])
            except Exception:
                rs_len = -1
            print(f"{log_prefix} 🎬 guided-mode recorded_steps_len={rs_len}")
        else:
            print(f"{log_prefix} 📝 freeform-mode (no recorded_steps in options)")


        expect = extract_success_expectations(task_text)
        print("🎯 Success expectations:", expect)
        primary_url = extract_primary_open_url(task_text)
        did_force_open = False

        max_rounds = int(_env("MAX_ROUNDS", "12"))
        max_actions_per_round = int(_env("MAX_ACTIONS_PER_ROUND", "6"))
        max_failures = int(_env("MAX_FAILURES", "10"))

        replay_enabled = False  # cache/replay disabled by design
        record_enabled = False  # cache/record disabled by design
        replayed_this_task = False
        recorded_plan: List[Dict[str, Any]] = []

        system = build_system_prompt(task_text, expect)

        prev_url: Optional[str] = None
        prev_title: Optional[str] = None

        
        # -----------------------------
        # Guided mode (recommended): execute recorded_steps sequentially,
        # but allow the LLM to choose the best TARGET strategy (label/role/placeholder/text)
        # for each click/type step.
        # -----------------------------
        recorded_steps = options.get("recorded_steps") or []
        guided_mode = isinstance(recorded_steps, list) and len(recorded_steps) > 0

        if guided_mode:
            print(f"{log_prefix} ▶ Starting guided execution (no cache/replay).")

            # Build a stable system prompt once
            system = build_system_prompt(task_text, expect)

            async def plan_for_step(step: Dict[str, Any], url: str, title: str, elements: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
                action = (step.get("action") or "").strip().lower()

                ctx = _compact_context(url, title, elements)

                if action == "click":
                    label = (step.get("name") or "").strip()
                    locators = step.get("locators") or []  # Get locators
                    
                    # NEW: Use fallback strategy if locators available
                    if locators:
                        # Sort locators by priority (same as type)
                        LOCATOR_PRIORITY = {
                            "test-id": 0, "id": 1, "name": 2, "href": 3,
                            "placeholder": 4, "alt": 5, "aria-label": 6, "title": 7,
                            "label": 8, "css": 9, "role": 10, "text": 11, "xpath": 12
                        }
                        sorted_locators = sorted(locators, key=lambda x: LOCATOR_PRIORITY.get(x.get("type"), 99))
                        
                        print(f"\n{log_prefix} 🔄 CLICK with fallback strategy:")
                        print(f"{log_prefix}    Target: '{label}'")
                        print(f"{log_prefix}    Trying {len(sorted_locators)} locators in priority order...")
                        
                        # Skip LLM entirely - send all locators to client for fallback
                        plan = [{
                            "action": "clickWithFallback",
                            "locators": sorted_locators
                        }]
                        return plan
                    
                    # OLD: No locators - use LLM (backward compatibility)
                    instruction = (
                        "You MUST execute the next recorded step EXACTLY.\n"
                        f"RECORDED STEP: click the element with VISIBLE TEXT '{label}'.\n"
                        "Return a JSON array with exactly ONE action: click.\n"
                        "Constraints:\n"
                        "- You MUST return exactly ONE click action targeting the intended element.\n"
                        "- For button/link text like '{label}', use by='text' (Example: {\"by\":\"text\",\"text\":\""+label+"\",\"exact\":false})\n"
                        "- Alternatively, use by='role' with name parameter (Example: {\"by\":\"role\",\"role\":\"button\",\"name\":\""+label+"\",\"exact\":false})\n"
                        "- ❌ NEVER use by='name' - it is NOT valid!\n"
                        "- DO NOT add scroll_page - Playwright auto-scrolls elements into view.\n"
                        "- DO NOT click any other element.\n"
                    )
                    sub_system = BASE_PROMPT + "\n\n" + ALLOWED_ACTIONS
                    user = instruction + "\n\n" + ctx
                    
                    # VERBOSE LOGGING
                    print(f"\n{log_prefix} 🤖 LLM PROMPT (CLICK):")
                    print(f"{log_prefix}    Target: '{label}'")
                    
                    out = await ollama_chat(sub_system, user)
                    
                    print(f"\n{log_prefix} 🤖 LLM RESPONSE:")
                    print(f"{log_prefix}    {out[:200]}...")
                    
                    plan = _extract_json_array(out)
                    print(f"{log_prefix} 📋 EXTRACTED PLAN: {plan}")
                    # filter out any non click/scroll
                    plan = [a for a in plan if (a.get("action") in ("scroll_page", "click"))]
                    return plan

                if action == "type":
                    field = (step.get("name") or "").strip()
                    value = (step.get("value") or "")
                    locators = step.get("locators") or []  # Get locators
                    
                    # NEW: Use fallback strategy if locators available
                    if locators:
                        # Sort locators by priority
                        LOCATOR_PRIORITY = {
                            "test-id": 0, "id": 1, "name": 2, "href": 3,
                            "placeholder": 4, "alt": 5, "aria-label": 6, "title": 7,
                            "label": 8, "css": 9, "role": 10, "text": 11, "xpath": 12
                        }
                        sorted_locators = sorted(locators, key=lambda x: LOCATOR_PRIORITY.get(x.get("type"), 99))
                        
                        print(f"\n{log_prefix} 🔄 TYPE with fallback strategy:")
                        print(f"{log_prefix}    Field: '{field}'")
                        print(f"{log_prefix}    Value: '{value}'")
                        print(f"{log_prefix}    Trying {len(sorted_locators)} locators in priority order...")
                        
                        # Skip LLM entirely - send all locators to client for fallback
                        plan = [{
                            "action": "typeWithFallback",
                            "locators": sorted_locators,
                            "text": value
                        }]
                        return plan
                    
                    # OLD: No locators - use LLM (backward compatibility)
                    instruction = (
                        "You MUST execute the next recorded step EXACTLY.\\n"
                        f"RECORDED STEP: type value '{value}' into the field labeled/name '{field}'.\\n"
                        "Return a JSON array of actions. Allowed: click, type.\\n"
                        "Constraints:\\n"
                        "- You MUST end with exactly ONE type action into the intended field.\\n"
                        "- Prefer by='placeholder' first (most forms use placeholder text).\\n"
                        "- If no placeholder, try by='label' or by='role' with name='textbox'.\\n"
                        "- You may add ONE click before type only if needed to focus the field.\\n"
                        "- DO NOT add scroll_page - Playwright auto-scrolls elements into view.\\n"
                        "- DO NOT type into the wrong field.\\n"
                    )
                    sub_system = BASE_PROMPT + "\n\n" + ALLOWED_ACTIONS
                    user = instruction + "\n\n" + ctx
                    
                    # VERBOSE LOGGING
                    print(f"\n{log_prefix} 🤖 LLM PROMPT (TYPE - no locators):")
                    print(f"{log_prefix}    Field: '{field}'")
                    print(f"{log_prefix}    Value: '{value}'")
                    print(f"{log_prefix}    Instruction: {instruction[:150]}...")
                    
                    out = await ollama_chat(sub_system, user)
                    
                    print(f"\n{log_prefix} 🤖 LLM RESPONSE:")
                    print(f"{log_prefix}    {out[:200]}...")
                    
                    plan = _extract_json_array(out)
                    print(f"{log_prefix} 📋 EXTRACTED PLAN: {plan}")
                    # filter out any non click/type/scroll
                    plan = [a for a in plan if (a.get("action") in ("scroll_page", "click", "type"))]
                    return plan

                return []

            # Execute each recorded step deterministically, using LLM only for click/type targeting
            for i, step in enumerate(recorded_steps, 1):
                st_action = (step.get("action") or "").strip().lower()
                print(f"{log_prefix} 🔹 Step {i}/{len(recorded_steps)}: {st_action} {step}")

                if st_action in ("open", "navigate"):
                    url_to = (step.get("url") or "").strip()
                    if not url_to:
                        raise RuntimeError(f"Recorded step missing url: {step}")
                    did_any_action = True
                    print(f"{log_prefix} 🌐 GOTO -> {url_to}")
                    await send_command("navigate", {"url": url_to})
                    try:
                        cur = await send_command("getCurrentUrl", {})
                        print(f"{log_prefix} ✅ LOADED -> {cur}")
                    except Exception:
                        pass
                    continue

                # Gather context for LLM targeting decisions
                url = await send_command("getCurrentUrl", {})
                title = await send_command("getTitle", {})
                elements = await send_command("getInteractiveElements", {}) or []

                if st_action in ("click", "type"):
                    planned = await plan_for_step(step, url, title, elements)
                    if not planned:
                        raise RuntimeError(f"LLM returned empty plan for step: {step}")

                    # Execute planned actions
                    for a in planned:
                        a = normalize_action(a)
                        await exec_action(a)

                    continue

                if st_action == "verify_text":
                    txt = (step.get("value") or "").strip()
                    if not txt:
                        raise RuntimeError(f"verify_text missing value: {step}")
                    print(f"{log_prefix} ✅ VERIFY_TEXT -> '{txt}'")
                    await send_command("waitForText", {"text": txt, "timeoutMs": 10000})
                    continue

                if st_action == "verify_visible":
                    txt = (step.get("value") or "").strip()
                    if not txt:
                        raise RuntimeError(f"verify_visible missing value: {step}")
                    print(f"{log_prefix} ✅ VERIFY_ELEMENT(text) -> '{txt}'")
                    await send_command("waitForText", {"text": txt, "timeoutMs": 10000})
                    continue

                if st_action == "extract":
                    # NEW: Handle extract actions during playback
                    await exec_action(step)
                    continue

                if st_action == "extract_table":
                    # NEW: Phase 8.3 - Handle table extraction during playback
                    await exec_action(step)
                    continue

                if st_action == "wait":
                    # NEW: Handle wait actions during playback
                    await exec_action(step)
                    continue

                print(f"{log_prefix} ⚠️ Skipping unknown recorded step: {step}")

            # After all recorded steps, complete task
            print(f"{log_prefix} 🎉 TASK COMPLETE (guided) success=True did_any_action={did_any_action}")
            await send_json({"type": "task-complete", "taskId": task_id, "success": True})
            return
        for round_idx in range(max_rounds):
            url = await send_command("getCurrentUrl", {})
            # ✅ Safety-net: ensure we are on the main URL before asking the model to "act"
            if primary_url and not did_force_open:
                if (not url) or (url == "about:blank") or (not url.startswith(primary_url)):
                    print(f"🧭 Forcing initial navigation to primary URL: {primary_url} (current: {url})")
                    await send_command("navigate", {"url": primary_url})
                    did_force_open = True
                    continue

            title = await send_command("getTitle", {})
            elements = await send_command("getInteractiveElements", {}) or []
            context = _compact_context(url, title, elements)

            strict = plan_requires_strict_success(task_text)
            url_ok = url_meets_expectations(url, expect.get("expected_url_substrings", []))
            text_ok = await verify_texts(send_command, expect.get("expected_texts", []))

            progress_ok = _has_progress(prev_url, prev_title, url, title)
            prev_url, prev_title = url, title


            explicit_verify = bool(expect.get("expected_url_substrings") or expect.get("expected_texts"))
            
            # Only allow early success if:
            # - user explicitly asked verification (Verify ... lines present)
            # - and we actually executed actions
            if strict and explicit_verify and did_any_action and (url_ok or text_ok):
                await send_json({"type": "task-complete", "taskId": task_id, "success": True})
                return


            # ---- LLM planning (single round for simplicity) ----
            user_prompt = build_subgoal_prompt("act", system) + "\n\n" + context + "\n\nTASK:\n" + task_text
            llm_out = await ollama_chat(system, user_prompt)
            plan = _extract_json_array(llm_out)

            if record_enabled:
                recorded_plan.extend(plan)

            print("📋 plan:", plan)

            actions_attempted += len(plan)

            for a in plan[:max_actions_per_round]:
                a = normalize_action(a)
                try:
                    await exec_action(a)
                except Exception as e:
                    failures += 1
                    last_errors.append(str(e))
                    if failures >= max_failures:
                        await send_json({
                            "type": "task-complete",
                            "taskId": task_id,
                            "success": False,
                            "error": f"Too many failures. Last error: {e}",
                        })
                        return

            if record_enabled and plan:
                cache_put_plan(url, task_text, recorded_plan)

        await send_json({
            "type": "task-complete",
            "taskId": task_id,
            "success": False,
            "error": f"Max rounds reached.\n{_format_last_errors(last_errors)}",
        })
    
    except TaskDone:
        return

    except Exception as e:
        await send_json({
            "type": "task-complete",
            "taskId": task_id or "unknown",
            "success": False,
            "error": str(e),
        })
        raise


async def main():
    """
    Entrypoint for the Local WebAI Server.
    
    Starts a WebSocket server that listens for connections from the 
    Playwright Browser Client. Initializes server configuration from 
    environment variables and sets up the event loop.
    """
    host = _env("HOST", "localhost")
    port = int(_env("PORT", "8765"))
    path = _env("WS_PATH", "/api")

    print("✅ Local AI WebSocket server (LLM): ws://localhost:8765/api")
    print(f"   Using Ollama: {_env('OLLAMA_URL', 'http://localhost:11434')}/api/chat | model: {_env('OLLAMA_MODEL', 'llama3.1')}")
    print(f"   Replay: {_env('REPLAY_ENABLED', '1')} Record: {_env('RECORD_ENABLED', '1')} Cache: plan_cache.json")
    print("   🔄 FALLBACK STRATEGY ENABLED (v2.0)  ← CODE UPDATED!")  # NEW: Verify code loaded

    async def handler(ws):
        req_path = get_request_path(ws)
        if not req_path.startswith(path):
            await ws.close()
            return
        await handle_client(ws)

    # IMPORTANT: max_size=None removes the 1MB limit that caused 1009 errors.
    async with websockets.serve(
        handler,
        host,
        port,
        max_size=None,      # <-- FIX
        max_queue=32,
        ping_interval=20,
        ping_timeout=60,
    ):
        await asyncio.Future()


if __name__ == "__main__":
    asyncio.run(main())
